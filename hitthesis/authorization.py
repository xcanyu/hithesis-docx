"""
授权声明页
本科和硕博共用，通过 thesis_type 区分内容
"""

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .ooxml_utils import set_font, set_first_line_indent, add_signature_line, set_outline_level


def add_authorization(doc, info: dict, add_to_toc=True, thesis_type="bachelor"):
    """添加授权声明

    Args:
        doc: python-docx Document 对象
        info: {"title": "...", "author": "...", "supervisor": "...", ...}
        add_to_toc: 是否在目录中添加条目
        thesis_type: "bachelor" | "master" | "doctor"
    """
    if thesis_type == "bachelor":
        _add_bachelor(doc, info, add_to_toc)
    else:
        _add_graduate(doc, info, add_to_toc)


def _add_bachelor(doc, info, add_to_toc):
    """本科毕业论文授权声明"""
    # 第一排：哈尔滨工业大学本科毕业论文（设计）
    para1 = doc.add_paragraph()
    pPr1 = para1._element.get_or_add_pPr()
    pb1 = OxmlElement('w:pageBreakBefore')
    pPr1.append(pb1)
    for tag in ['w:jc', 'w:spacing', 'w:keepNext', 'w:keepLines', 'w:snapToGrid']:
        old = pPr1.find(qn(tag))
        if old is not None:
            pPr1.remove(old)
    jc1 = OxmlElement('w:jc')
    jc1.set(qn('w:val'), 'center')
    pPr1.append(jc1)
    spacing1 = OxmlElement('w:spacing')
    spacing1.set(qn('w:before'), '196')
    spacing1.set(qn('w:after'), '0')
    pPr1.append(spacing1)
    pPr1.append(OxmlElement('w:keepNext'))
    pPr1.append(OxmlElement('w:keepLines'))
    sg1 = OxmlElement('w:snapToGrid')
    sg1.set(qn('w:val'), '0')
    pPr1.append(sg1)
    run1 = para1.add_run('哈尔滨工业大学本科毕业论文（设计）')
    set_font(run1, '黑体', 18, False)

    # 第二排：原创性声明和使用权限
    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.line_spacing = Pt(20.5)
    para2.paragraph_format.space_before = Pt(3)
    para2.paragraph_format.space_after = Pt(12)
    run2 = para2.add_run('原创性声明和使用权限')
    set_font(run2, '黑体', 18, False)

    # 第三排：本科毕业论文（设计）原创性声明
    para3_title = doc.add_paragraph()
    para3_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para3_title.paragraph_format.line_spacing = Pt(20.5)
    para3_title.paragraph_format.space_before = Pt(23)
    para3_title.paragraph_format.space_after = Pt(14)
    run3t = para3_title.add_run('本科毕业论文（设计）原创性声明')
    set_font(run3t, '黑体', 15, False)

    # 正文内容
    title = info.get("title", "（论文题目）")
    content_text = (
        f'本人郑重声明：此处所提交的本科毕业论文（设计）《{title}》，是本人在导师指导下，'
        f'在哈尔滨工业大学攻读学士学位期间独立进行研究工作所取得的成果，且毕业论文（设计）'
        f'中除已标注引用文献的部分外不包含他人完成或已发表的研究成果。对本毕业论文（设计）'
        f'的研究工作做出重要贡献的个人和集体，均已在文中以明确方式注明。'
        f'本毕业论文（设计）对使用AI工具的情况进行了明确标注。'
    )
    para3 = doc.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para3.paragraph_format.line_spacing = Pt(20.5)
    para3.paragraph_format.space_before = Pt(0)
    para3.paragraph_format.space_after = Pt(0)
    set_first_line_indent(para3, 480)
    run3 = para3.add_run(content_text)
    set_font(run3, '宋体', 12)
    rPr3 = run3._element.get_or_add_rPr()
    sp3 = OxmlElement('w:spacing')
    sp3.set(qn('w:val'), '10')
    rPr3.append(sp3)

    # 签名行（原创性声明后）
    para_sign = doc.add_paragraph()
    para_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_sign.paragraph_format.line_spacing = Pt(20.5)
    para_sign.paragraph_format.space_before = Pt(28)
    add_signature_line(para_sign, '作者签名：', 114)

    # 本科毕业论文（设计）使用权限
    para4 = doc.add_paragraph()
    para4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para4.paragraph_format.line_spacing = Pt(24)
    para4.paragraph_format.space_before = Pt(49)
    para4.paragraph_format.space_after = Pt(10)
    sg = OxmlElement('w:snapToGrid')
    sg.set(qn('w:val'), '0')
    para4._element.get_or_add_pPr().append(sg)
    run4 = para4.add_run('本科毕业论文（设计）使用权限')
    set_font(run4, '黑体', 15, False)

    # 使用权限正文
    para5 = doc.add_paragraph()
    para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para5.paragraph_format.line_spacing = Pt(20.5)
    para5.paragraph_format.space_before = Pt(0)
    para5.paragraph_format.space_after = Pt(0)
    set_first_line_indent(para5, 480)
    run5 = para5.add_run(
        '本科毕业论文（设计）是本科生在哈尔滨工业大学攻读学士学位期间完成的成果，'
        '知识产权归属哈尔滨工业大学。本科毕业论文（设计）的使用权限如下：'
    )
    set_font(run5, '宋体', 12)

    line_data = [
        ('（1）学校可以采用影印、缩印或其他复制手段保存本科生上交的毕业论文', WD_ALIGN_PARAGRAPH.LEFT, 6, True),
        ('（设计），并向有关部门报送本科毕业论文（设计）；（2）根据需要，学校可', WD_ALIGN_PARAGRAPH.LEFT, 6, False),
        ('本科毕业论文（设计）部分或全部内容编入有关数据库进行检索和提供相应阅', WD_ALIGN_PARAGRAPH.JUSTIFY, 10, False),
        ('览服务；（3）本科生毕业后发表与此毕业论文（设计）研究成果相关的学术论', WD_ALIGN_PARAGRAPH.LEFT, 6, False),
        ('文和其他成果时，应征得导师同意，且第一署名单位为哈尔滨工业大学。', WD_ALIGN_PARAGRAPH.LEFT, 6, False),
    ]
    for text, align, spacing_val, is_first in line_data:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = Pt(20.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if is_first:
            set_first_line_indent(p, 480)
        run = p.add_run(text)
        set_font(run, '宋体', 12)
        rPr = run._element.get_or_add_rPr()
        sp = rPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            rPr.append(sp)
        sp.set(qn('w:val'), str(spacing_val))

    # 保密段落
    para7 = doc.add_paragraph()
    para7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para7.paragraph_format.line_spacing = Pt(20.5)
    para7.paragraph_format.space_before = Pt(0)
    para7.paragraph_format.space_after = Pt(0)
    set_first_line_indent(para7, 480)
    run7 = para7.add_run('保密论文在保密期内遵守有关保密规定，解密后适用于此使用权限规定。')
    set_font(run7, '宋体', 12)
    rPr7 = run7._element.get_or_add_rPr()
    sp7 = OxmlElement('w:spacing')
    sp7.set(qn('w:val'), '10')
    rPr7.append(sp7)

    # 本人知悉段落
    para8 = doc.add_paragraph()
    para8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para8.paragraph_format.line_spacing = Pt(20.5)
    para8.paragraph_format.space_before = Pt(0)
    para8.paragraph_format.space_after = Pt(0)
    set_first_line_indent(para8, 480)
    run8 = para8.add_run('本人知悉本科毕业论文（设计）的使用权限，并将遵守有关规定。')
    set_font(run8, '宋体', 12)
    rPr8 = run8._element.get_or_add_rPr()
    sp8 = OxmlElement('w:spacing')
    sp8.set(qn('w:val'), '10')
    rPr8.append(sp8)

    # 作者签名（使用权限后）
    para_auth_sign = doc.add_paragraph()
    para_auth_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_auth_sign.paragraph_format.line_spacing = Pt(20.5)
    para_auth_sign.paragraph_format.space_before = Pt(34)
    para_auth_sign.paragraph_format.space_after = Pt(0)
    add_signature_line(para_auth_sign, '作者签名：', 114)

    # 导师签名
    para_sup_sign = doc.add_paragraph()
    para_sup_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_sup_sign.paragraph_format.line_spacing = Pt(20.5)
    para_sup_sign.paragraph_format.space_before = Pt(16)
    para_sup_sign.paragraph_format.space_after = Pt(0)
    add_signature_line(para_sup_sign, '导师签名：', 114)

    # 目录条目
    if add_to_toc:
        hidden_para = doc.add_paragraph()
        set_outline_level(hidden_para, 0)
        hidden_para.paragraph_format.space_before = Pt(0)
        hidden_para.paragraph_format.space_after = Pt(0)
        hidden_para.paragraph_format.line_spacing = 1
        hidden_run = hidden_para.add_run(
            '哈尔滨工业大学本科毕业论文（设计）原创性声明和使用权限'
        )
        hidden_run.font.size = Pt(1)
        hidden_run.font.color.rgb = RGBColor(255, 255, 255)


def _add_graduate(doc, info, add_to_toc):
    """硕博学位论文授权声明"""
    # 第一排：哈尔滨工业大学学位论文原创性声明和使用权限
    para1 = doc.add_paragraph()
    pPr1 = para1._element.get_or_add_pPr()
    pb1 = OxmlElement('w:pageBreakBefore')
    pPr1.append(pb1)
    for tag in ['w:jc', 'w:spacing', 'w:keepNext', 'w:keepLines', 'w:snapToGrid']:
        old = pPr1.find(qn(tag))
        if old is not None:
            pPr1.remove(old)
    jc1 = OxmlElement('w:jc')
    jc1.set(qn('w:val'), 'center')
    pPr1.append(jc1)
    spacing1 = OxmlElement('w:spacing')
    spacing1.set(qn('w:before'), '310')
    spacing1.set(qn('w:after'), '0')
    pPr1.append(spacing1)
    pPr1.append(OxmlElement('w:keepNext'))
    pPr1.append(OxmlElement('w:keepLines'))
    sg1 = OxmlElement('w:snapToGrid')
    sg1.set(qn('w:val'), '0')
    pPr1.append(sg1)
    run1 = para1.add_run('哈尔滨工业大学学位论文原创性声明和使用权限')
    set_font(run1, '黑体', 18, False)

    # 第二排：学位论文原创性声明（距第一排1.6cm）
    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.line_spacing = Pt(20.5)
    para2.paragraph_format.space_before = Pt(45.36)
    para2.paragraph_format.space_after = Pt(12)
    run2 = para2.add_run('学位论文原创性声明')
    set_font(run2, '黑体', 15, False)

    # 正文内容
    title = info.get("title", "（论文题目）")
    content_text = (
        f'本人郑重声明：此处所提交的学位论文《{title}》，是本人在导师指导下，'
        f'在哈尔滨工业大学攻读学位期间独立进行研究工作所取得的成果，且论文'
        f'中除已标注引用文献的部分外不包含他人完成或已发表的研究成果。对本论文'
        f'的研究工作做出重要贡献的个人和集体，均已在文中以明确方式注明。'
    )
    para3 = doc.add_paragraph()
    para3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para3.paragraph_format.line_spacing = Pt(20.5)
    para3.paragraph_format.space_before = Pt(0)
    para3.paragraph_format.space_after = Pt(0)
    set_first_line_indent(para3, 480)
    run3 = para3.add_run(content_text)
    set_font(run3, '宋体', 12.5)

    # 签名行
    para_sign = doc.add_paragraph()
    para_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_sign.paragraph_format.line_spacing = Pt(20.5)
    para_sign.paragraph_format.space_before = Pt(14)
    add_signature_line(para_sign, '作者签名：', 114)

    # 学位论文使用权限
    para4 = doc.add_paragraph()
    para4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para4.paragraph_format.line_spacing = Pt(24)
    para4.paragraph_format.space_before = Pt(56.7)
    para4.paragraph_format.space_after = Pt(10)
    sg = OxmlElement('w:snapToGrid')
    sg.set(qn('w:val'), '0')
    para4._element.get_or_add_pPr().append(sg)
    run4 = para4.add_run('学位论文使用权限')
    set_font(run4, '黑体', 15, False)

    # 使用权限正文
    para5 = doc.add_paragraph()
    para5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para5.paragraph_format.line_spacing = Pt(20.5)
    para5.paragraph_format.space_before = Pt(0)
    para5.paragraph_format.space_after = Pt(0)
    set_first_line_indent(para5, 480)
    run5 = para5.add_run(
        '学位论文是研究生在哈尔滨工业大学攻读学位期间完成的成果，'
        '知识产权归属哈尔滨工业大学。学位论文的使用权限如下：'
    )
    set_font(run5, '宋体', 12.5)

    # 条款正文
    line_data = [
        ('（1）学校可以采用影印、缩印或其他复制手段保存研究生上交的学位论文，', WD_ALIGN_PARAGRAPH.LEFT, -1, True),
        ('并向国家图书馆报送学位论文；（2）学校可以将学位论文部分或全部内容编入', WD_ALIGN_PARAGRAPH.JUSTIFY, 5, False),
        ('有关数据库进行检索和提供相应阅览服务；（3）研究生毕业后发表与此学位论', WD_ALIGN_PARAGRAPH.JUSTIFY, 5, False),
        ('文研究成果相关的学术论文和其他成果时，应征得导师同意，且第一署名单位', WD_ALIGN_PARAGRAPH.JUSTIFY, 9, False),
        ('为哈尔滨工业大学。', WD_ALIGN_PARAGRAPH.JUSTIFY, 0, False),
    ]
    for text, align, spacing_val, is_first in line_data:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = Pt(20.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if is_first:
            set_first_line_indent(p, 480)
        run = p.add_run(text)
        set_font(run, '宋体', 12)
        rPr = run._element.get_or_add_rPr()
        sp = rPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            rPr.append(sp)
        sp.set(qn('w:val'), str(spacing_val))

    # 保密段落
    para7 = doc.add_paragraph()
    para7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para7.paragraph_format.line_spacing = Pt(20.5)
    para7.paragraph_format.space_before = Pt(0)
    para7.paragraph_format.space_after = Pt(0)
    set_first_line_indent(para7, 480)
    run7 = para7.add_run('保密论文在保密期内遵守有关保密规定，解密后适用于此使用权限规定。')
    set_font(run7, '宋体', 12)
    rPr7 = run7._element.get_or_add_rPr()
    sp7 = OxmlElement('w:spacing')
    sp7.set(qn('w:val'), '10')
    rPr7.append(sp7)

    # 本人知悉段落
    para8 = doc.add_paragraph()
    para8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para8.paragraph_format.line_spacing = Pt(20.5)
    para8.paragraph_format.space_before = Pt(0)
    para8.paragraph_format.space_after = Pt(0)
    set_first_line_indent(para8, 480)
    run8 = para8.add_run('本人知悉学位论文的使用权限，并将遵守有关规定。')
    set_font(run8, '宋体', 12.5)

    # 作者签名
    para_auth_sign = doc.add_paragraph()
    para_auth_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_auth_sign.paragraph_format.line_spacing = Pt(20.5)
    para_auth_sign.paragraph_format.space_before = Pt(23)
    para_auth_sign.paragraph_format.space_after = Pt(0)
    add_signature_line(para_auth_sign, '作者签名：', 114)

    # 导师签名
    para_sup_sign = doc.add_paragraph()
    para_sup_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_sup_sign.paragraph_format.line_spacing = Pt(20.5)
    para_sup_sign.paragraph_format.space_before = Pt(17)
    para_sup_sign.paragraph_format.space_after = Pt(0)
    add_signature_line(para_sup_sign, '导师签名：', 114)

    # 目录条目
    if add_to_toc:
        hidden_para = doc.add_paragraph()
        set_outline_level(hidden_para, 0)
        hidden_para.paragraph_format.space_before = Pt(0)
        hidden_para.paragraph_format.space_after = Pt(0)
        hidden_para.paragraph_format.line_spacing = 1
        hidden_run = hidden_para.add_run(
            '哈尔滨工业大学学位论文原创性声明和使用权限'
        )
        hidden_run.font.size = Pt(1)
        hidden_run.font.color.rgb = RGBColor(255, 255, 255)
