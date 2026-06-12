"""断字集成测试：验证 compile 后的 docx 包含正确的断字设置

生成的测试文档可在 Word 中打开，手动验证断字效果：
1. 打开 output/hyphenation_test.docx
2. 确保 Word 开启自动断字：File → Options → Proofing → Automatically hyphenate this document when typing
3. 查看参考文献节的英文长词是否在行尾按音节断字
"""

import os
import zipfile
import tempfile
import pytest
from hitthesis import Thesis
from hitthesis.docx_postproc import enable_auto_hyphenation, set_paragraph_lang_in_references


@pytest.fixture
def output_dir():
    """创建输出目录"""
    out = os.path.join(os.path.dirname(__file__), "..", "output")
    os.makedirs(out, exist_ok=True)
    return out


def _read_settings_xml(path):
    """读取 docx 中的 settings.xml"""
    with zipfile.ZipFile(path) as z:
        return z.read("word/settings.xml").decode("utf-8")


def _read_document_xml(path):
    """读取 docx 中的 document.xml"""
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


class TestHyphenationIntegration:
    """断字集成测试"""

    def test_compile_has_hyphenation_settings(self, output_dir):
        """compile 后的 docx 应包含自动断字设置"""
        doc = Thesis(type="bachelor", campus="harbin")
        doc.set_info(
            title="断字测试",
            author="测试作者",
            supervisor="测试导师"
        )
        doc.add_paragraph("测试内容")

        output_path = os.path.join(output_dir, "hyphenation_test.docx")
        doc.compile(output_path)

        # 验证 settings.xml 包含断字设置
        settings_xml = _read_settings_xml(output_path)
        assert "autoHyphenation" in settings_xml, "缺少 autoHyphenation 设置"
        assert "hyphenationZone" in settings_xml, "缺少 hyphenationZone 设置"
        assert 'val="200"' in settings_xml, "hyphenationZone 应为 200"

    def test_compile_has_lang_in_references(self, output_dir):
        """compile 后的 docx 参考文献节应有语言设置"""
        doc = Thesis(type="bachelor", campus="harbin")
        doc.set_info(
            title="断字测试",
            author="测试作者",
            supervisor="测试导师"
        )
        doc.add_paragraph("参考文献")
        doc.add_paragraph(
            "GROSS W A. Thermal constriction in gas-lubricated bearings[J]. "
            "Journal of Tribology, 1973, 95(1): 1-8."
        )

        output_path = os.path.join(output_dir, "hyphenation_refs_test.docx")
        doc.compile(output_path)

        # 验证参考文献节有语言设置
        doc_xml = _read_document_xml(output_path)
        assert 'w:val="en-US"' in doc_xml, "参考文献节应设置 lang=en-US"
        assert 'w:eastAsia="zh-CN"' in doc_xml, "参考文献节应设置 eastAsia=zh-CN"

    def test_generate_manual_verification_doc(self, output_dir):
        """生成用于手动验证断字效果的文档

        此测试生成一个包含长英文词的文档，供用户在 Word 中打开验证。
        """
        doc = Thesis(type="bachelor", campus="harbin")
        doc.set_info(
            title="英文断字效果验证",
            author="测试作者",
            supervisor="测试导师"
        )

        doc.add_paragraph("参考文献")
        # 故意用长英文词，测试断字效果
        refs = [
            "SMITH J, JOHNSON B. Internationalization and characterization of lubrication mechanisms in tribological systems[J]. Journal of Tribology, 2023, 145(3): 1-15.",
            "WILLIAMS R, BROWN T. Electromagnetohydrodynamic phenomena in microfluidic applications[J]. Physics of Fluids, 2022, 34(8): 082001.",
            "CHEN L, WANG H. Characterization of microstructural evolution in nanocrystalline materials during severe plastic deformation[J]. Acta Materialia, 2021, 195: 112-125.",
            "GARCIA M, LOPEZ P. Counterrevolutionary and antidisestablishmentarianism perspectives in modern sociolinguistics[J]. Language, 2020, 96(2): 45-67.",
            "ZHANG Y, LI W. Electroencephalographic and magnetohydrodynamic analysis of neurological disorders[J]. NeuroImage, 2019, 189: 312-328.",
        ]
        for ref in refs:
            doc.add_paragraph(ref)

        output_path = os.path.join(output_dir, "hyphenation_manual_test.docx")
        doc.compile(output_path)

        # 验证文件生成成功
        assert os.path.exists(output_path)
        assert os.path.getsize(output_path) > 0

        # 打印提示
        print(f"\n{'='*60}")
        print(f"手动验证文档已生成: {output_path}")
        print(f"请在 Word 中打开，查看参考文献节的英文长词是否断字")
        print(f"{'='*60}")


class TestHyphenationPostProc:
    """断字后处理函数单元测试"""

    def test_enable_auto_hyphenation_idempotent(self):
        """重复调用 enable_auto_hyphenation 应幂等"""
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            from docx import Document
            doc = Document()
            doc.save(path)

            # 多次调用
            enable_auto_hyphenation(path)
            enable_auto_hyphenation(path)
            enable_auto_hyphenation(path)

            settings = _read_settings_xml(path)
            assert settings.count("<w:autoHyphenation") == 1
            assert settings.count("<w:hyphenationZone") == 1

    def test_set_paragraph_lang_idempotent(self):
        """重复调用 set_paragraph_lang_in_references 应幂等"""
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "test.docx")
            from docx import Document
            doc = Document()
            doc.add_paragraph("参考文献")
            doc.add_paragraph("SMITH J. Test article[J]. Journal, 2023.")
            doc.save(path)

            # 多次调用
            set_paragraph_lang_in_references(path)
            first = _read_document_xml(path).count("<w:lang ")
            set_paragraph_lang_in_references(path)
            second = _read_document_xml(path).count("<w:lang ")

            assert first == second


if __name__ == "__main__":
    # 直接运行时生成手动验证文档
    import sys
    out = os.path.join(os.path.dirname(__file__), "..", "output")
    os.makedirs(out, exist_ok=True)

    doc = Thesis(type="bachelor", campus="harbin")
    doc.set_info(
        title="英文断字效果验证",
        author="测试作者",
        supervisor="测试导师"
    )

    doc.add_paragraph("参考文献")
    refs = [
        "SMITH J, JOHNSON B. Internationalization and characterization of lubrication mechanisms in tribological systems[J]. Journal of Tribology, 2023, 145(3): 1-15.",
        "WILLIAMS R, BROWN T. Electromagnetohydrodynamic phenomena in microfluidic applications[J]. Physics of Fluids, 2022, 34(8): 082001.",
        "CHEN L, WANG H. Characterization of microstructural evolution in nanocrystalline materials during severe plastic deformation[J]. Acta Materialia, 2021, 195: 112-125.",
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    output_path = os.path.join(out, "hyphenation_manual_test.docx")
    doc.compile(output_path)
    print(f"已生成: {output_path}")
