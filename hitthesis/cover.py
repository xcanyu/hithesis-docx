"""
封面 + 本科生第二封面 (cover2)
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from .config import UNIVERSITY_NAME
from .ooxml_utils import (
    set_font, estimate_text_width, set_cell_vertical_alignment,
    disable_snap_to_grid, add_spacer, make_tbl_borders_none,
    suppress_auto_hyphenation,
)


def add_cover(doc, info: dict, thesis_type: str, english_title: str = None):
    """添加封面

    Args:
        doc: python-docx Document 对象
        info: {"title": "...", "author": "...", "date": "...", ...}
        thesis_type: 当前仅实现 "bachelor"，"master"/"doctor" 预留
        english_title: 英文标题，默认使用预置文本
    """
    cover_first_line = {
        "bachelor": "本科毕业论文（设计）",
        "master": "硕士学位论文",
        "doctor": "博士学位论文",
    }.get(thesis_type, "学位论文")

    # 封面间距校准值（基于 2025 年官方本科 Word 模板 VBA 实测 + 直尺验证，误差 < 2mm）
    # 详见 dev-notes.md "封面间距校准"

    # 第一行
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = Pt(42)
    para.paragraph_format.space_before = Pt(34)
    para.paragraph_format.space_after = Pt(52)
    run = para.add_run(cover_first_line)
    set_font(run, "宋体", 24, True)
    suppress_auto_hyphenation(para)

    # 标题 - 黑体二号（22pt），居中
    para_title = doc.add_paragraph()
    para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_title.paragraph_format.line_spacing = Pt(22)
    para_title.paragraph_format.space_before = Pt(10)
    para_title.paragraph_format.space_after = Pt(44)
    run_title = para_title.add_run(info.get("title", ""))
    set_font(run_title, "黑体", 22, False)
    suppress_auto_hyphenation(para_title)

    # 英文标题
    if english_title is None:
        english_title = info.get("english_title", "")
    if english_title:
        text_area_width = 425.2 * 0.9
        W_22 = estimate_text_width(english_title, 22)
        W_18 = estimate_text_width(english_title, 18)
        if W_22 <= text_area_width:
            en_font_size = 22
        elif W_18 > 425.2:
            en_font_size = 18
        else:
            en_font_size = 22
    else:
        en_font_size = 22

    if english_title:
        para_en = doc.add_paragraph()
        para_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_en.paragraph_format.line_spacing = Pt(25)
        para_en.paragraph_format.space_before = Pt(17.4)
        para_en.paragraph_format.space_after = Pt(67.2)
        run_en = para_en.add_run(english_title)
        set_font(run_en, "Times New Roman", en_font_size, True)
        suppress_auto_hyphenation(para_en)

    # 作者姓名
    para_author = doc.add_paragraph()
    para_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_author.paragraph_format.line_spacing = Pt(22)
    para_author.paragraph_format.space_before = Pt(0)
    para_author.paragraph_format.space_after = Pt(139.2)
    run_author = para_author.add_run(info.get("author", ""))
    set_font(run_author, "宋体", 18, True)
    suppress_auto_hyphenation(para_author)

    # 哈尔滨工业大学
    para_univ = doc.add_paragraph()
    para_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_univ.paragraph_format.line_spacing = Pt(22)
    para_univ.paragraph_format.space_before = Pt(0)
    para_univ.paragraph_format.space_after = Pt(9.7)
    run_univ = para_univ.add_run("哈尔滨工业大学")
    set_font(run_univ, "楷体", 18, True)
    suppress_auto_hyphenation(para_univ)

    # 日期
    para_date = doc.add_paragraph()
    para_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_date.paragraph_format.line_spacing = Pt(22)
    para_date.paragraph_format.space_before = Pt(0)
    para_date.paragraph_format.space_after = Pt(14.25)
    date_str = info.get("date", "")
    run_date = para_date.add_run(date_str)
    set_font(run_date, "宋体", 18, True)
    suppress_auto_hyphenation(para_date)

    # 本科生第二封面
    if thesis_type == "bachelor":
        add_cover2_bachelor(doc, info)


def add_cover2_bachelor(doc, info: dict):
    """添加第二个封面页的信息表格（本科生模板）

    纯段落定位（无 wp:anchor 锚点），避免 Word 拒绝打开文档：
    - 密级：右对齐 + 右缩进0.3cm，页面顶部
    - "本科毕业论文（设计）"：居中，距页上边7.1cm
    - 论文标题：居中，距页上边9.55cm
    - 信息表格：tblInd 水平定位
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 分节符（新页面）
    section = doc.add_section()
    section.start_type = WD_SECTION_START.NEW_PAGE

    # 1. 密级行：左侧勾选项 + 右侧密级
    para_secret = doc.add_paragraph()
    para_secret.paragraph_format.space_before = Pt(0)
    para_secret.paragraph_format.space_after = Pt(0)
    para_secret.paragraph_format.line_spacing = Pt(14)
    disable_snap_to_grid(para_secret)
    suppress_auto_hyphenation(para_secret)
    # 清除列表格式
    pPr = para_secret._element.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)
    outlineLvl = pPr.find(qn('w:outlineLvl'))
    if outlineLvl is not None:
        pPr.remove(outlineLvl)
    # 左侧：勾选项（TNR空格分隔）
    run_left1 = para_secret.add_run("☐毕业论文")
    set_font(run_left1, "宋体", 12, False)
    run_space1 = para_secret.add_run("    ")  # TNR四格约12.5pt
    set_font(run_space1, "Times New Roman", 12, False)
    run_left2 = para_secret.add_run("☐毕业设计")
    set_font(run_left2, "宋体", 12, False)
    # 用TNR空格推到右侧（224.6pt ≈ 75个TNR空格）
    run_space = para_secret.add_run(" " * 77)
    set_font(run_space, "Times New Roman", 12, False)
    run_right = para_secret.add_run("密级：公开")
    set_font(run_right, "宋体", 12, False)

    # 2. 空白
    add_spacer(doc, 58.3)

    # 3. "本科毕业论文（设计）"
    para_label = doc.add_paragraph()
    para_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_label.paragraph_format.space_before = Pt(0)
    para_label.paragraph_format.space_after = Pt(0)
    para_label.paragraph_format.line_spacing = Pt(22)
    disable_snap_to_grid(para_label)
    suppress_auto_hyphenation(para_label)
    run = para_label.add_run("本科毕业论文（设计）")
    set_font(run, "宋体", 18, True)

    # 4. 空白
    add_spacer(doc, 45.1)

    # 5. 论文标题
    thesis_title = info.get("title", "")
    para_thesis = doc.add_paragraph()
    para_thesis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_thesis.paragraph_format.space_before = Pt(0)
    para_thesis.paragraph_format.space_after = Pt(0)
    para_thesis.paragraph_format.line_spacing = Pt(28)
    disable_snap_to_grid(para_thesis)
    suppress_auto_hyphenation(para_thesis)
    run = para_thesis.add_run(thesis_title)
    set_font(run, "黑体", 22, False)

    # 6. 空白到表格
    add_spacer(doc, 159.2)

    # 7. 学生信息表格
    left_indent_twips = int((171 - 3.0 / 2.54 * 72) * 20)

    degree_label = {
        "bachelor": "本科生",
        "master": "硕士生",
        "doctor": "博士生",
    }.get(info.get("type", "bachelor"), "本科生")

    info_items = [
        (degree_label, info.get("author", "")),
        ("学号", info.get("student_id", "")),
        ("指导教师", info.get("supervisor", "")),
        ("专业", info.get("subject", "")),
        ("学院", info.get("affil", "")),
        ("答辩日期", info.get("date", "")),
        ("学校", UNIVERSITY_NAME),
    ]

    table = doc.add_table(rows=7, cols=3)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(left_indent_twips))
    tblInd.set(qn('w:type'), 'dxa')
    old_ind = tblPr.find(qn('w:tblInd'))
    if old_ind is not None:
        tblPr.remove(old_ind)
    tblPr.insert(0, tblInd)

    col_widths = [int(3.19 * 567), int(0.53 * 567), int(5.2 * 567)]
    total_width = sum(col_widths)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total_width))
    tblW.set(qn('w:type'), 'dxa')
    old_tblW = tblPr.find(qn('w:tblW'))
    if old_tblW is not None:
        tblPr.remove(old_tblW)
    tblPr.append(tblW)

    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblPr.append(make_tbl_borders_none())

    row_height = 583

    for row_idx, (label, value) in enumerate(info_items):
        row = table.rows[row_idx]
        trPr = row._tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row._tr.insert(0, trPr)
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(row_height))
        trHeight.set(qn('w:hRule'), 'exact')
        old_height = trPr.find(qn('w:trHeight'))
        if old_height is not None:
            trPr.remove(old_height)
        trPr.append(trHeight)

        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths[col_idx]))
            tcW.set(qn('w:type'), 'dxa')
            old_tcW = tcPr.find(qn('w:tcW'))
            if old_tcW is not None:
                tcPr.remove(old_tcW)
            tcPr.append(tcW)

            set_cell_vertical_alignment(cell, "center")

            cellMar = tcPr.find(qn('w:tcMar'))
            if cellMar is not None:
                tcPr.remove(cellMar)
            cellMar = OxmlElement('w:tcMar')
            for margin in ('top', 'right', 'bottom', 'left'):
                m = OxmlElement(f'w:{margin}')
                m.set(qn('w:w'), '0')
                m.set(qn('w:type'), 'dxa')
                cellMar.append(m)
            tcPr.append(cellMar)

            para = cell.paragraphs[0]
            para._element.clear()

            if col_idx == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            pPr = para._element.get_or_add_pPr()
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:before'), '152')
            spacing.set(qn('w:after'), '152')
            spacing.set(qn('w:line'), '360')
            spacing.set(qn('w:lineRule'), 'exact')

            if col_idx == 0:
                run = para.add_run(label)
                set_font(run, "黑体", 14, False)
            elif col_idx == 1:
                run = para.add_run("：")
                set_font(run, "宋体", 14, False)
            else:
                run = para.add_run(value)
                set_font(run, "宋体", 14, False)
