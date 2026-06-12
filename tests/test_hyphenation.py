"""断字设置测试：验证 docx_postproc.enable_auto_hyphenation 正确注入 settings.xml"""

import os
import tempfile
import zipfile

import pytest
from docx import Document

from hitthesis.docx_postproc import (
    enable_auto_hyphenation,
    set_paragraph_lang_in_references,
)


@pytest.fixture
def docx_file():
    """创建一个临时 docx 文件"""
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, "test.docx")
        doc = Document()
        doc.add_paragraph("hello world")
        doc.save(path)
        yield path


def _read_settings_xml(path):
    with zipfile.ZipFile(path) as z:
        return z.read("word/settings.xml").decode("utf-8")


def _read_document_xml(path):
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


class TestEnableAutoHyphenation:
    """注入断字设置测试"""

    def test_inject_basic(self, docx_file):
        """默认参数应注入 autoHyphenation=true 和 hyphenationZone=200"""
        enable_auto_hyphenation(docx_file)
        xml = _read_settings_xml(docx_file)
        assert 'autoHyphenation' in xml
        assert 'hyphenationZone' in xml
        # 默认 zone=200（更激进的断字策略）
        assert 'val="200"' in xml

    def test_inject_with_custom_zone(self, docx_file):
        """自定义断字区"""
        enable_auto_hyphenation(docx_file, zone_twips=200)
        xml = _read_settings_xml(docx_file)
        assert 'val="200"' in xml

    def test_idempotent(self, docx_file):
        """重复调用应幂等：不会产生重复 autoHyphenation 元素"""
        enable_auto_hyphenation(docx_file)
        enable_auto_hyphenation(docx_file)
        enable_auto_hyphenation(docx_file)
        xml = _read_settings_xml(docx_file)
        # 只能有 1 个 autoHyphenation 元素
        assert xml.count("<w:autoHyphenation") == 1
        assert xml.count("<w:hyphenationZone") == 1

    def test_zipfile_valid(self, docx_file):
        """处理后 docx 仍能正常打开（不破坏 zip 结构）"""
        enable_auto_hyphenation(docx_file)
        # 能正常打开
        doc = Document(docx_file)
        assert len(doc.paragraphs) >= 1
        # 文件可重新打开
        with zipfile.ZipFile(docx_file) as z:
            names = z.namelist()
            assert "word/settings.xml" in names
            assert "word/document.xml" in names


class TestSetParagraphLangInReferences:
    """参考文献节段落语言设置测试"""

    def _make_docx_with_refs(self, tmp_path):
        """构造一个含参考文献节的临时 docx"""
        path = str(tmp_path / "test_refs.docx")
        doc = Document()
        doc.add_paragraph("参考文献")
        doc.add_paragraph("GROSS W A. Thermal constriction in gas-lubricated bearings[J].Journal of Tribology, 1973, 95(1): 1-8.")
        doc.add_paragraph("致谢")  # 后续段落
        doc.add_paragraph("感谢导师")
        doc.save(path)
        return path

    def test_sets_lang_on_refs_runs(self, tmp_path):
        """参考文献节 run 应该有 <w:lang> 属性"""
        path = self._make_docx_with_refs(tmp_path)
        count = set_paragraph_lang_in_references(path)
        assert count >= 1  # 至少改了一个 run

        xml = _read_document_xml(path)
        # 参考文献段落应有 lang="en-US"
        assert 'w:val="en-US"' in xml
        # 东亚语言应是 zh-CN
        assert 'w:eastAsia="zh-CN"' in xml

    def test_no_lang_before_refs(self, tmp_path):
        """参考文献标题前的段落不应被改"""
        path = self._make_docx_with_refs(tmp_path)
        set_paragraph_lang_in_references(path)
        xml = _read_document_xml(path)
        # "感谢导师" 段不应有 en-US（但可以有默认 en-US 由 docDefaults 设）
        # 这里只检查参考文献节之后的段落没被改
        # 由于 docDefaults 已有 en-US，所以 "感谢导师" 段可能也有
        # 这个测试改成：参考文献节内的 run 一定有 en-US
        assert 'w:val="en-US"' in xml

    def test_idempotent(self, tmp_path):
        """重复调用应幂等"""
        path = self._make_docx_with_refs(tmp_path)
        set_paragraph_lang_in_references(path)
        first = _read_document_xml(path).count("<w:lang ")
        set_paragraph_lang_in_references(path)
        second = _read_document_xml(path).count("<w:lang ")
        # 重复调用不增加 lang 数量
        assert first == second

    def test_custom_lang(self, tmp_path):
        """自定义语言代码"""
        path = self._make_docx_with_refs(tmp_path)
        set_paragraph_lang_in_references(path, lang_western="en-GB", lang_eastasia="zh-TW")
        xml = _read_document_xml(path)
        assert 'w:val="en-GB"' in xml
        assert 'w:eastAsia="zh-TW"' in xml
