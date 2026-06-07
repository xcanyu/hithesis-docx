"""
文档元素：表格、图片、子图
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .ooxml_utils import set_font, set_cell_vertical_alignment, add_caption_with_bookmark


class Table:
    """表格元素 - 三线表"""

    def __init__(self, thesis, rows, cols, caption=None, label=None, ref=None):
        self.thesis = thesis  # Thesis 对象引用
        self.doc = thesis.doc  # Document 对象
        self.rows = rows
        self.cols = cols
        self.caption = caption
        self.label = label  # 编号，如 "1-1"
        self.ref = ref  # 引用标签
        self._table = None

    def _clear_cell_margins(self, cell):
        """清除单元格的内边距（单位：twips），确保内容完全居中"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        cellMar = tcPr.find(qn('w:tcMar'))
        if cellMar is not None:
            tcPr.remove(cellMar)
        cellMar = OxmlElement('w:tcMar')
        for margin in ('top', 'right', 'bottom', 'left'):
            m = OxmlElement(f'w:{margin}')
            m.set(qn('w:w'), '0')
            m.set(qn('w:type'), 'dxa')
            cellMar.append(m)
        tcPr.append(cellMar)

    def insert(self):
        """插入表格"""
        if self.caption:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = Pt(14)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            text = f"表{self.label}　{self.caption}" if self.label else self.caption
            add_caption_with_bookmark(para, self.ref, text)

        self._table = self.doc.add_table(rows=self.rows, cols=self.cols)
        self._table.alignment = WD_TABLE_ALIGNMENT.CENTER

        self._apply_three_line_style()

        for row in self._table.rows:
            for cell in row.cells:
                set_cell_vertical_alignment(cell, "center")
                self._clear_cell_margins(cell)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.line_spacing = Pt(14)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        set_font(run, "宋体", 10.5, False)

        return self

    def _apply_three_line_style(self):
        r"""应用三线表样式（booktabs风格）
        顶线/底线 = 1.5磅 = w:sz=12，中线 = 1磅 = w:sz=8
        """
        if not self._table:
            return

        tbl = self._table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        def make_tc_border(tag, sz, color='000000'):
            b = OxmlElement(f'w:{tag}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), str(sz))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), color)
            return b

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tblBorders.append(b)
        old_tbl_borders = tblPr.find(qn('w:tblBorders'))
        if old_tbl_borders is not None:
            tblPr.remove(old_tbl_borders)
        tblPr.append(tblBorders)

        row_count = len(self._table.rows)
        for row_idx, row in enumerate(self._table.rows):
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)

                old_tcBorders = tcPr.find(qn('w:tcBorders'))
                if old_tcBorders is not None:
                    tcPr.remove(old_tcBorders)

                tcBorders = OxmlElement('w:tcBorders')
                if row_idx == 0:
                    tcBorders.append(make_tc_border('top', 12))
                    tcBorders.append(make_tc_border('bottom', 8))
                elif row_idx == row_count - 1:
                    tcBorders.append(make_tc_border('bottom', 12))

                if len(tcBorders) > 0:
                    tcPr.append(tcBorders)

    def set_cell(self, row, col, text, bold=False):
        """设置单元格内容（支持 $...$ 渲染）"""
        if self._table and row < len(self._table.rows) and col < len(self._table.columns):
            cell = self._table.rows[row].cells[col]
            para = cell.paragraphs[0]
            para.clear()
            if '$' in text and '$' in text[1:]:
                self.thesis._add_text_with_superscripts(para, text)
            else:
                run = para.add_run(str(text))
                set_font(run, "宋体", 10.5, bold)


class Figure:
    """图片元素"""

    def __init__(self, thesis, image_path=None, caption=None, width=None, label=None, ref=None):
        self.thesis = thesis
        self.doc = thesis.doc
        self.image_path = image_path
        self.caption = caption
        self.width = width
        self.label = label
        self.ref = ref

    def insert(self):
        """插入图片"""
        if self.image_path:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(self.image_path, width=self.width or Cm(10))

        if self.caption:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = Pt(14)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            text = f"图{self.label}　{self.caption}" if self.label else self.caption
            add_caption_with_bookmark(para, self.ref, text)

        return self


class SubFigure:
    """子图元素（支持并排或网格布局，每个子图有 (a)(b)(c) 标号）"""

    def __init__(self, thesis, subfigures, caption=None, label=None, ref=None, cols=None):
        self.thesis = thesis
        self.doc = thesis.doc
        self.subfigures = subfigures
        self.caption = caption
        self.label = label
        self.ref = ref
        self.cols = cols

    def insert(self):
        """插入子图"""
        n = len(self.subfigures)

        if self.cols:
            cols = self.cols
        else:
            cols = n
        rows = (n + cols - 1) // cols

        table = self.doc.add_table(rows=rows * 2, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tblBorders.append(b)
        old_borders = tblPr.find(qn('w:tblBorders'))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblPr.append(tblBorders)

        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        old_tblW = tblPr.find(qn('w:tblW'))
        if old_tblW is not None:
            tblPr.remove(old_tblW)
        tblPr.append(tblW)

        letters = 'abcdefghijklmnopqrstuvwxyz'

        for i, (image_path, subcaption) in enumerate(self.subfigures):
            row_idx = (i // cols) * 2
            col_idx = i % cols

            cell = table.cell(row_idx, col_idx)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            if image_path:
                run = para.add_run()
                run.add_picture(image_path, width=Cm(10 / cols))

            cell_cap = table.cell(row_idx + 1, col_idx)
            para_cap = cell_cap.paragraphs[0]
            para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_cap.paragraph_format.space_before = Pt(2)
            para_cap.paragraph_format.space_after = Pt(0)
            letter = letters[i] if i < len(letters) else str(i + 1)
            text = f"({letter}) {subcaption}" if subcaption else f"({letter})"
            run = para_cap.add_run(text)
            set_font(run, "宋体", 10.5)

        if self.caption:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = Pt(14)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            text = f"图{self.label}　{self.caption}" if self.label else self.caption
            add_caption_with_bookmark(para, self.ref, text)

        return self
