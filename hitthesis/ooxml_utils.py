"""
OOXML 原子操作工具集
统一管理所有底层 Word OOXML 操作，消除 document.py 中的重复 XML 模式
"""

import re
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# 字体
# ============================================================================

def set_font(run, cn_font, size, bold=False):
    """设置中英文字体：中文用指定字体，英文固定 Times New Roman

    中文引号""''等属于 Unicode 通用标点区段，Word 用 hAnsi 字体渲染。
    若 run 文本含 CJK 或中文标点，则 hAnsi 也用中文字体，避免引号变英文样式。
    """
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), "Times New Roman")
    # 若文本含中文字符或中文标点，hAnsi 也用中文字体
    text = run.text or ""
    if _has_cjk_or_cn_punct(text):
        rFonts.set(qn('w:hAnsi'), cn_font)
    else:
        rFonts.set(qn('w:hAnsi'), "Times New Roman")


def _has_cjk_or_cn_punct(text: str) -> bool:
    """检测文本是否含 CJK 字符或中文标点（含中文引号等）"""
    for ch in text:
        cp = ord(ch)
        # CJK 统一表意文字
        if 0x4E00 <= cp <= 0x9FFF:
            return True
        # 中文引号和标点：“ ” ‘ ’ — … 、 。
        if cp in (0x201C, 0x201D, 0x2018, 0x2019, 0x2014, 0x2026, 0x3001, 0x3002):
            return True
        # 全角空格、中文顿号等
        if 0x3000 <= cp <= 0x303F:
            return True
    return False


def estimate_text_width(text, font_size_pt):
    """估算英文文本宽度（pt），按字符类型分别估算宽度"""
    width = 0
    for ch in text:
        if 'A' <= ch <= 'Z' or 'a' <= ch <= 'z':
            width += 0.52 * font_size_pt
        elif ch == ' ':
            width += 0.3 * font_size_pt
        else:
            width += 0.55 * font_size_pt
    return width


# ============================================================================
# 单元格
# ============================================================================

def set_cell_vertical_alignment(cell, alignment="center"):
    """设置单元格垂直对齐方式（center / top / bottom）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for vAlign in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(vAlign)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    tcPr.append(vAlign)


def clear_cell_margins(cell):
    """清除单元格所有内边距"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    cellMar = tcPr.find(qn('w:tcMar'))
    if cellMar is not None:
        tcPr.remove(cellMar)
    cellMar = OxmlElement('w:tcMar')
    for margin in ('top', 'right', 'bottom', 'left'):
        m = OxmlElement(f'w:{margin}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        cellMar.append(m)
    tcPr.append(cellMar)


# ============================================================================
# 边框
# ============================================================================

def make_border_element(tag, val='single', sz=4, space=0, color='000000'):
    """创建单个边框元素"""
    b = OxmlElement(f'w:{tag}')
    b.set(qn('w:val'), val)
    b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), str(space))
    b.set(qn('w:color'), color)
    return b


def make_tbl_borders_none():
    """创建无边框的 tblBorders 元素（表格级清除所有边框）"""
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tblBorders.append(make_border_element(name, 'none', 0, 0, 'auto'))
    return tblBorders


def set_paragraph_border(para, border_position, sz, space, color):
    """设置段落边框（例如下边框）
    border_position: 'top' | 'left' | 'bottom' | 'right'
    sz: 边框粗细（1/8pt 单位，18 = 2.25pt）
    space: 边框与文字的距离（pt）
    """
    pPr = para._element.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    old = pBdr.find(qn(f'w:{border_position}'))
    if old is not None:
        pBdr.remove(old)
    pBdr.append(make_border_element(border_position, 'single', sz, space, color))


def set_cell_bottom_border(cell, sz, color='000000'):
    """给单元格设置仅下边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    old = tcBorders.find(qn('w:bottom'))
    if old is not None:
        tcBorders.remove(old)
    tcBorders.append(make_border_element('bottom', 'single', sz, 0, color))


def remove_table_borders(table):
    """移除表格所有边框，替换为 nil"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(make_tbl_borders_none())


def apply_tc_borders(cell, parts: dict):
    """按指定部分设置单元格边框，如 {'top': 12, 'bottom': 8}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for position, sz in parts.items():
        tcBorders.append(make_border_element(position, 'single', sz, 0, '000000'))
    tcPr.append(tcBorders)


# ============================================================================
# 段落格式
# ============================================================================

def set_outline_level(para, level):
    """设置段落大纲级别（供 Word 目录识别，0=Heading1, 1=Heading2...）"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level))
    pPr.append(outlineLvl)


def add_page_break_before(para):
    """w:pageBreakBefore — 段落从新页开始"""
    pPr = para._element.get_or_add_pPr()
    pPr.append(OxmlElement('w:pageBreakBefore'))


def apply_heading_style(para, after_pt=None):
    """一级标题样式：居中、段后18pt、1.2倍行距、与下段同页、不对齐到网格
    after_pt: 段后间距，默认取 SPACING["heading_after"]
    """
    from .config import SPACING
    if after_pt is None:
        after_pt = int(SPACING["heading_after"] / 20)
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    for tag in ['w:jc', 'w:spacing', 'w:keepNext', 'w:keepLines', 'w:snapToGrid']:
        old = pPr.find(qn(tag))
        if old is not None:
            pPr.remove(old)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(SPACING["heading_before"]))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    spacing.set(qn('w:line'), str(SPACING["heading_line"]))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    pPr.append(OxmlElement('w:keepNext'))
    pPr.append(OxmlElement('w:keepLines'))
    sg = OxmlElement('w:snapToGrid')
    sg.set(qn('w:val'), '0')
    pPr.append(sg)


def add_heading_properties(para):
    """为节标题添加 keepNext/keepLines/snapToGrid（轻量版）"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    for tag in ['keepNext', 'keepLines']:
        if pPr.find(qn(f'w:{tag}')) is None:
            pPr.append(OxmlElement(f'w:{tag}'))
    if pPr.find(qn('w:snapToGrid')) is None:
        sg = OxmlElement('w:snapToGrid')
        sg.set(qn('w:val'), '0')
        pPr.append(sg)


def set_first_line_indent(para, twips):
    """直接设置 OOXML w:firstLine 缩进（twips）"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), str(twips))


def set_hanging_indent(para, twips):
    """设置悬挂缩进：首行顶格，续行缩进 twips（2字符=480twips）
    等价于 w:left=twips + w:firstLine=-twips"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), str(twips))
    ind.set(qn('w:firstLine'), str(-twips))


def disable_snap_to_grid(para):
    """禁用段落对齐网格"""
    pPr = para._element.get_or_add_pPr()
    snap = OxmlElement('w:snapToGrid')
    snap.set(qn('w:val'), '0')
    pPr.append(snap)


def suppress_auto_hyphenation(para):
    """禁用段落自动断字（用于封面等不应断字的段落）"""
    pPr = para._element.get_or_add_pPr()
    suppress = pPr.find(qn('w:suppressAutoHyphens'))
    if suppress is None:
        suppress = OxmlElement('w:suppressAutoHyphens')
        pPr.append(suppress)
    suppress.set(qn('w:val'), 'true')


def add_spacer(doc, height_pt):
    """添加精确高度的空白段落（space_before + snapToGrid=0）"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(height_pt)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(1)
    disable_snap_to_grid(para)
    run = para.add_run('​')
    set_font(run, "宋体", 1)
    return para


# ============================================================================
# 标题段落创建
# ============================================================================

def make_heading_para(doc, text, font_cn='黑体', font_size=18, bold=False):
    """创建一级标题段落（章/附录共用），复用 apply_heading_style"""
    para = doc.add_paragraph()
    add_page_break_before(para)
    apply_heading_style(para)
    set_outline_level(para, 0)
    run = para.add_run(text)
    set_font(run, font_cn, font_size, bold)
    return para


# ============================================================================
# 书签
# ============================================================================

_BOOKMARK_ID_COUNTER = [1000]  # 全局递增 ID（从 1000 起步，避免与系统预定义冲突）


def _next_bookmark_id() -> int:
    """获取下一个唯一 bookmark ID。"""
    _BOOKMARK_ID_COUNTER[0] += 1
    return _BOOKMARK_ID_COUNTER[0]


def add_caption_with_bookmark(para, ref, text, font_cn="宋体", font_size=10.5, reference_db=None):
    """添加带书签的题注文本（图/表/代码/子图共用）

    支持 caption 中的 [ref:key] 引用标记，渲染为上标文献序号。
    """
    if ref:
        bm_id = str(_next_bookmark_id())
        bm_start = OxmlElement('w:bookmarkStart')
        bm_start.set(qn('w:id'), bm_id)
        bm_start.set(qn('w:name'), f"cite_{ref}")
        para._element.append(bm_start)
        _add_caption_runs(para, text, font_cn, font_size, reference_db)
        bm_end = OxmlElement('w:bookmarkEnd')
        bm_end.set(qn('w:id'), bm_id)
        bm_end.set(qn('w:name'), f"cite_{ref}")
        para._element.append(bm_end)
    else:
        _add_caption_runs(para, text, font_cn, font_size, reference_db)


def _add_caption_runs(para, text, font_cn="宋体", font_size=10.5, reference_db=None):
    """渲染题注文本，将 [ref:key] 解析为上标超链接引用。"""
    parts = re.split(r'(\[ref:[a-zA-Z0-9_:.-]+\])', text)
    for part in parts:
        if not part:
            continue
        m = re.match(r'\[ref:([a-zA-Z0-9_:.-]+)\]', part)
        if m:
            key = m.group(1)
            if reference_db:
                num = reference_db.cite(key)
                lb = para.add_run('[')
                set_font(lb, "Times New Roman", font_size)
                lb.font.superscript = True
                add_ref_hyperlink(para, str(num), f'ref_{num}')
                rb = para.add_run(']')
                set_font(rb, "Times New Roman", font_size)
                rb.font.superscript = True
            else:
                run = para.add_run(f'[ref:{key}]')
                set_font(run, "Times New Roman", font_size)
                run.font.superscript = True
        else:
            run = para.add_run(part)
            set_font(run, font_cn, font_size)


def add_bookmark(para_element, name):
    """在段落 XML 元素中添加书签（起止标记）。

    使用全局递增 ID 避免与文档其他 bookmark 冲突。
    """
    bm_id = str(_next_bookmark_id())
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), bm_id)
    bm_start.set(qn('w:name'), name)
    para_element.append(bm_start)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), bm_id)
    bm_end.set(qn('w:name'), name)
    para_element.append(bm_end)


def add_bookmark_to_paragraph(para, bookmark_name: str):
    """为文档段落添加书签。"""
    bm_id = str(_next_bookmark_id())
    bookmarkStart = OxmlElement('w:bookmarkStart')
    bookmarkStart.set(qn('w:id'), bm_id)
    bookmarkStart.set(qn('w:name'), bookmark_name)
    para._p.insert(0, bookmarkStart)
    bookmarkEnd = OxmlElement('w:bookmarkEnd')
    bookmarkEnd.set(qn('w:id'), bm_id)
    bookmarkEnd.set(qn('w:name'), bookmark_name)
    para._p.append(bookmarkEnd)


# ============================================================================
# 超链接
# ============================================================================

def add_ref_hyperlink(para, text, anchor):
    """添加参考文献超链接（上标样式）"""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    hyperlink.set(qn('w:history'), '1')
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')
    rPr.append(szCs)
    vert = OxmlElement('w:vertAlign')
    vert.set(qn('w:val'), 'superscript')
    rPr.append(vert)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


def add_hyperlink_run(para, text, anchor):
    """向段落添加超链接（黑色无下划线，字体与周围一致）"""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    hyperlink.set(qn('w:history'), '1')
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'none')
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    para._element.append(hyperlink)


# ============================================================================
# 签名行
# ============================================================================

def add_signature_line(para, label, width_after=114):
    """右对齐签名行：标签 + 间距 + '日期：' + '年' + '月' + '日'"""
    def _seg(text):
        r = para.add_run(text)
        set_font(r, '宋体', 12)
        return r

    def _space(w_pt, n=4):
        run = para.add_run(' ' * n)
        set_font(run, '宋体', 12)
        rPr = run._element.get_or_add_rPr()
        sp = OxmlElement('w:spacing')
        tracking = int((w_pt * 20 - n * 80) / max(n - 1, 1))
        if tracking > 0:
            sp.set(qn('w:val'), str(tracking))
            rPr.append(sp)
        return run

    _seg(label)
    _space(width_after, 6)
    _seg('日期：')
    _space(31.7, 4)
    _seg('年')
    _space(18.875, 3)
    _seg('月')
    _space(18.875, 3)
    _seg('日')


# ============================================================================
# 引用文本渲染
# ============================================================================

def render_normal_text_with_superscripts(para, text: str):
    """渲染普通文本中的 [数字] 上标"""
    import re
    sub_parts = re.split(r'\[(\d+(?:-\d+)?)\]', text)
    for j, sub_part in enumerate(sub_parts):
        if j % 2 == 1 and re.match(r'^\d+(?:-\d+)?$', sub_part):
            run = para.add_run(f'[{sub_part}]')
            run.font.superscript = True
            set_font(run, "Times New Roman", 12)
        elif sub_part:
            run = para.add_run(sub_part)
            set_font(run, "宋体", 12)


def add_ref_runs(para, keys: list, reference_db=None, min_run_length: int = 2):
    """渲染参考文献引用，连续序号自动合并为区间。

    输出形态（GB/T 7714 风格）：
      单个：              [5]
      连续区间：          [1-3]   （当区间长度 > min_run_length）
      独立+区间混合：     [1,3-5]
      全部独立：          [1,2,3] （区间长度 <= min_run_length 时不合并）

    Args:
        para: python-docx Paragraph
        keys: 引用 key 列表（按出现顺序）
        reference_db: ReferenceDB 实例；若为 None 则原样输出 [ref:key] 占位
        min_run_length: 至少几个连续序号才合并，默认为 2。
                        设 1 = 全部独立显示；设 2 = 连续 2 个以上合并
    """
    if not keys:
        return
    if reference_db is None:
        for key in keys:
            run = para.add_run(f'[ref:{key}]')
            set_font(run, "Times New Roman", 12)
            run.font.superscript = True
        return
    nums = [reference_db.cite(key) for key in keys]
    nums = sorted(set(nums))
    items = []
    i = 0
    while i < len(nums):
        start = nums[i]
        end = start
        j = i + 1
        while j < len(nums) and nums[j] == end + 1:
            end = nums[j]
            j += 1
        length = end - start + 1
        if length > min_run_length:
            items.append((start, end))
        else:
            for num in range(start, end + 1):
                items.append(num)
        i = j
    left_bracket = para.add_run('[')
    set_font(left_bracket, "Times New Roman", 12)
    left_bracket.font.superscript = True
    for idx, item in enumerate(items):
        if isinstance(item, int):
            add_ref_hyperlink(para, str(item), f'ref_{item}')
        else:
            start, end = item
            add_ref_hyperlink(para, f'{start}-{end}', f'ref_{start}')
        if idx != len(items) - 1:
            comma = para.add_run(',')
            set_font(comma, "Times New Roman", 12)
            comma.font.superscript = True
    right_bracket = para.add_run(']')
    set_font(right_bracket, "Times New Roman", 12)
    right_bracket.font.superscript = True


def group_consecutive_refs(text: str):
    """将文本中所有 [ref:xxx] 标记按"是否直接相邻"分组（中间 0 字符算一组）。

    这是 P 规则的核心函数：
      "同一处引用" = 两个 [ref:...] 之间没有任何字符
      "分别引用"   = 中间有任何字符（标点、空白、汉字都算独立）

    Args:
        text: 含 [ref:xxx] 标记的文本
    Returns:
        list[list[str]]: 每组是直接相邻的 [ref:...] key 列表
                         例如 [[A,B,C], [D], [E]] 表示 [A][B][C][D]...文本...[E]
    """
    import re
    pattern = re.compile(r'\[ref:([a-zA-Z0-9_:.\-]+)\]')
    matches = list(pattern.finditer(text))
    if not matches:
        return []

    groups = []
    current = [matches[0].group(1)]
    for prev, curr in zip(matches, matches[1:]):
        if curr.start() == prev.end():
            current.append(curr.group(1))
        else:
            groups.append(current)
            current = [curr.group(1)]
    groups.append(current)
    return groups
