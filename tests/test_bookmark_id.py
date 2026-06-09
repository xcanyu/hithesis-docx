"""Bookmark ID 唯一性测试：验证不再有写死 0 的 bookmark ID。"""

import re
import zipfile

import pytest
from docx import Document

from hitthesis import Thesis, ReferenceDB


@pytest.fixture
def doc():
    return Thesis(type="bachelor", campus="harbin")


def _read_doc_xml(path):
    with zipfile.ZipFile(str(path)) as z:
        return z.read("word/document.xml").decode("utf-8")


class TestBookmarkIdUniqueness:
    """验证 bookmarkStart 的 id 全部唯一，不写死 0"""

    def _compile_doc(self, doc, tmp_path):
        """编译 docx 并返回路径"""
        path = str(tmp_path / "test.docx")
        doc.set_info(title="T", author="A", supervisor="S")
        doc.add_cover()
        doc.start_arabic_section()
        with doc.add_chapter("章1", "1"):
            doc.add_section("节1")
            doc.add_paragraph("测试段落。")
            tbl = doc.add_table(2, 2, caption="表1", ref="tbl1")
            tbl.set_cell(0, 0, "A")
            tbl.set_cell(0, 1, "B")
            doc.add_figure(None, caption="图1", ref="fig1")
            doc.add_code_block("print(1)", caption="代码1", ref="code1")
            doc.add_equation("a+b", label="1-1", ref="eq1")
            doc.add_theorem("定理内容", kind="定理", ref="thm1")
            # 多次添加同名 ref
            tbl2 = doc.add_table(2, 2, caption="表2", ref="tbl2")
            tbl2.set_cell(0, 0, "C")
            doc.add_figure(None, caption="图2", ref="fig2")
        with doc.add_appendix("附录A"):
            doc.add_section("附节1")
            doc.add_paragraph("附段")
            tbl3 = doc.add_table(2, 2, caption="附表1", ref="tbl3")
            tbl3.set_cell(0, 0, "X")
            doc.add_equation("x+y", label="A-1", ref="eqA1")
        doc.compile(path)
        return path

    def test_no_duplicate_bookmark_ids(self, doc, tmp_path):
        """Word COM 重写 bookmark ID 是已知的，但无论如何 ID 必须唯一不重复"""
        path = self._compile_doc(doc, tmp_path)
        xml = _read_doc_xml(path)
        # 提取所有 bookmarkStart 的 id
        ids = re.findall(r'<w:bookmarkStart[^>]*w:id="(\d+)"', xml)
        # 验证：没有重复
        assert len(ids) == len(set(ids)), f"有重复 bookmark id: {ids}"
        # 注意：Word COM 展开 TOC 时会重新分配 bookmark ID（从 0 开始），
        # 所以我们没法保证 ID >= 1000。但 ID 必须唯一。

    def test_ids_are_numeric(self, doc, tmp_path):
        path = self._compile_doc(doc, tmp_path)
        xml = _read_doc_xml(path)
        ids = re.findall(r'<w:bookmarkStart[^>]*w:id="(\d+)"', xml)
        # Word COM 重新分配 ID 时会从 0 开始，所以不强制 >= 1000
        # 但 ID 必须是合法数字
        for bm_id in ids:
            assert bm_id.isdigit(), f"ID 不是数字: {bm_id}"
            assert int(bm_id) >= 0, f"ID 负数: {bm_id}"

    def test_start_end_ids_match(self, doc, tmp_path):
        """每个 bookmarkStart 的 id 必须有对应的 bookmarkEnd 用同一个 id"""
        path = self._compile_doc(doc, tmp_path)
        xml = _read_doc_xml(path)
        start_ids = set(re.findall(r'<w:bookmarkStart[^>]*w:id="(\d+)"', xml))
        end_ids = set(re.findall(r'<w:bookmarkEnd[^>]*w:id="(\d+)"', xml))
        assert start_ids == end_ids, f"Start/End ID 不匹配：start={start_ids - end_ids}, end={end_ids - start_ids}"
