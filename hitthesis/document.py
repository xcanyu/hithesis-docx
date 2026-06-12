"""
哈工大学位论文Word文档构建器
严格遵循《哈尔滨工业大学研究生学位论文撰写规范》
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from .config import PAGE, UNIVERSITY_NAME, SPACING
from .ooxml_utils import (
    set_font,
    set_paragraph_border,
    set_outline_level, add_page_break_before, apply_heading_style,
    add_heading_properties, set_first_line_indent,
    add_bookmark, add_bookmark_to_paragraph,
    add_ref_hyperlink, add_hyperlink_run, add_ref_runs, group_consecutive_refs,
    render_normal_text_with_superscripts, make_heading_para,
    disable_snap_to_grid,
    set_hanging_indent,
    add_caption_with_bookmark,
)
from .latex_render import add_latex_to_paragraph
from .cover import add_cover as _add_cover
from .authorization import add_authorization as _add_authorization
from .compile import compile_document
from .elements import Table, Figure, SubFigure
from .contexts import ChapterContext, AppendixContext


class Thesis:
    """学位论文文档构建器（当前仅实现本科/本部，硕博/其他校区预留）"""

    def __init__(self, type="bachelor", campus="harbin", header_text=None):
        self.doc = Document()
        self.type = type
        self.campus = campus
        if header_text is not None:
            self._header_text = header_text
        elif type == "bachelor":
            self._header_text = "哈尔滨工业大学本科毕业论文（设计）"
        else:
            self._header_text = UNIVERSITY_NAME
        self.chapter_number = 0  # 章节计数器
        self.section_counter = [0, 0, 0]  # [一级, 二级, 三级] 计数器
        self.appendix_number = 0  # 附录计数器
        self._eq_counter = 0  # 章节公式计数器
        self._cite_registry = {}  # 交叉引用注册表 {标签名: (类型, 编号)}
        self._pending_cites = []  # 待解析引用 [(para_element, run_element, tag), ...]
        self._reference_db = None  # ReferenceDB 实例
        self._footnotes = []  # 脚注收集 [(global_id, display_num, text), ...]
        self._footnote_counter = 0  # 脚注全局递增序号
        self._footnote_counter = 0  # 脚注全局递增序号（用于①②③）
        self._tab_counter = 0  # 表格计数器
        self._fig_counter = 0  # 图片计数器
        self._code_counter = 0  # 代码块计数器
        self._thm_counter = 0  # 定理论引理计数器
        self.info = {
            "title": "",
            "author": "",
            "supervisor": "",
            "subject": "",
            "affil": "",
            "date": "",
            "keywords": [],
        }
        self._setup_page()
        self._toc_blank_line = False  # 目录第一章前是否空一行
        self._in_appendix = False  # 是否在附录节（控制编号加"附"前缀）
        # 封面：无页眉、无页脚、无页码（不调用 _setup_header_footer）

    def set_reference_db(self, bib: 'ReferenceDB'):
        """设置文献数据库，用于 [ref:key] 引用"""
        self._reference_db = bib

    def _setup_page(self):
        """设置页面"""
        section = self.doc.sections[0]
        section.page_width = Cm(PAGE["width"])
        section.page_height = Cm(PAGE["height"])
        section.left_margin = Cm(PAGE["left"])
        section.right_margin = Cm(PAGE["right"])
        section.top_margin = Cm(PAGE["top"])
        section.bottom_margin = Cm(PAGE["bottom"])
        section.header_distance = Cm(PAGE["header"])
        section.footer_distance = Cm(PAGE["footer"])
        # 封面节：断开页眉页脚链接（保持空白）
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

    def _setup_page_numbering(self):
        """设置封面无页码（移除第一节的页码类型）"""
        section = self.doc.sections[0]
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            sectPr.remove(pgNumType)

    def _setup_header_footer_for_section(self, section):
        """为指定 section 设置页眉和页脚"""
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Pt, Cm

        # 断开链接
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        # ========== 页眉 ==========
        header = section.header
        # 清除所有现有的段落（确保从空白开始）
        for para in header.paragraphs:
            p = para._element
            p.getparent().remove(p)

        if self._header_text is False or self._header_text is None:
            return  # 不添加页眉

        # 1. 学校名称段落（粗线下边框，粗线与文字间距由space参数控制）
        para_title = header.add_paragraph()
        para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para_title.add_run(self._header_text)
        set_font(run, "宋体", 9)
        para_title.paragraph_format.space_after = Pt(0)
        set_paragraph_border(para_title, 'bottom', sz=18, space=0, color='000000')  # 粗线2.25pt

        # 2. 间距段落（控制粗线与细线间距 0.75pt）
        gap_para = header.add_paragraph()
        gap_para.paragraph_format.space_before = Pt(0)
        gap_para.paragraph_format.space_after = Pt(0)
        gap_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        gap_para.paragraph_format.line_spacing = Pt(0.75)

        # 3. 细线段落（细线下边框）
        para_thin = header.add_paragraph()
        para_thin.paragraph_format.space_before = Pt(0)
        para_thin.paragraph_format.space_after = Pt(0)
        para_thin.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para_thin.paragraph_format.line_spacing = Pt(1)
        run_thin = para_thin.add_run(' ')
        set_paragraph_border(para_thin, 'bottom', sz=6, space=0, color='000000')  # 细线0.75pt

        # ========== 页脚（页码：— PAGE —） ==========
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 前破折号
        run1 = footer_para.add_run("- ")
        run1.font.name = "Times New Roman"
        run1.font.size = Pt(9)

        # PAGE 域
        run2 = footer_para.add_run()
        rPr2 = run2._element.get_or_add_rPr()
        rFonts2 = OxmlElement('w:rFonts')
        rFonts2.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts2.set(qn('w:ascii'), 'Times New Roman')
        rFonts2.set(qn('w:hAnsi'), 'Times New Roman')
        rPr2.append(rFonts2)
        sz2 = OxmlElement('w:sz')
        sz2.set(qn('w:val'), '18')
        rPr2.append(sz2)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run2._r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        run2._r.append(instrText)
        run3 = footer_para.add_run()
        rPr3 = run3._element.get_or_add_rPr()
        rFonts3 = OxmlElement('w:rFonts')
        rFonts3.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts3.set(qn('w:ascii'), 'Times New Roman')
        rFonts3.set(qn('w:hAnsi'), 'Times New Roman')
        rPr3.append(rFonts3)
        sz3 = OxmlElement('w:sz')
        sz3.set(qn('w:val'), '18')
        rPr3.append(sz3)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)

        # 后破折号
        run4 = footer_para.add_run(" -")
        run4.font.name = "Times New Roman"
        run4.font.size = Pt(9)




    def _add_text_with_superscripts(self, para, text):
        """解析文本中的特殊标记并渲染为 OOXML。

        支持四种标记（可混用）：
          [ref:key]       → 参考文献引用（上标，连续的直接相邻自动合并为 [1-3]）
          [cite:tag]      → 交叉引用（超链接跳转到图表/公式）
          [数字]           → 手动上标编号
          $LaTeX$         → LaTeX 公式/化学式

        处理方式：先收集所有标记位置，按顺序分段渲染。
        连续的 [ref:]（中间 0 字符 = 直接相邻）先收集再批量输出（P 规则合并）。
        """
        import re
        last_end = 0
        positions = []

        # 收集所有需要特殊处理的位置
        for m in re.finditer(r'\[ref:([a-zA-Z0-9_:.-]+)\]', text):
            positions.append((m.start(), m.end(), 'ref', m.group(1)))
        for m in re.finditer(r'\[cite:([a-zA-Z0-9_:.-]+)\]', text):
            positions.append((m.start(), m.end(), 'cite', m.group(1)))
        for m in re.finditer(r'\[(\d+(?:-\d+)?)\]', text):
            positions.append((m.start(), m.end(), 'num', m.group(1)))
        for m in re.finditer(r'\$([^$]+)\$', text):
            positions.append((m.start(), m.end(), 'latex', m.group(1)))

        positions.sort(key=lambda x: (x[0], x[1]))

        # 按 P 规则分组连续的 [ref:]：直接相邻的合并为一组调用 add_ref_runs
        i = 0
        while i < len(positions):
            start, end, ptype, content2 = positions[i]
            if start > last_end:
                render_normal_text_with_superscripts(para, text[last_end:start])

            if ptype == 'ref':
                # 收集所有直接相邻的 [ref:] 直到遇到非 ref 或非相邻
                group = [content2]
                group_end = end
                j = i + 1
                while j < len(positions):
                    next_start, next_end, next_type, next_content = positions[j]
                    if next_type != 'ref' or next_start != group_end:
                        break
                    group.append(next_content)
                    group_end = next_end
                    j += 1
                add_ref_runs(para, group, self._reference_db)
                last_end = group_end
                i = j
                continue

            if ptype == 'cite':
                self._add_cite_runs(para, content2)
            elif ptype == 'num':
                run = para.add_run(f'[{content2}]')
                run.font.superscript = True
                set_font(run, "Times New Roman", 12)
            elif ptype == 'latex':
                add_latex_to_paragraph(para, content2, set_font)
            last_end = end
            i += 1

        if last_end < len(text):
            render_normal_text_with_superscripts(para, text[last_end:])

    def _add_cite_runs(self, para, cite_str):
        """向段落添加交叉引用超链接（支持延迟解析）"""
        import re
        tags = re.split(r'\]\s*,\s*\[', cite_str)
        for idx, tag in enumerate(tags):
            tag = tag.strip()
            if tag in self._cite_registry:
                ref_type, ref_num = self._cite_registry[tag]
                bookmark_name = f"cite_{tag}"
                if ref_type == "table":
                    display = f" {ref_num}"
                elif ref_type == "equation":
                    display = f" ({ref_num})"
                elif ref_type == "figure":
                    display = f" {ref_num}"
                else:
                    display = f" {ref_num}"
                add_hyperlink_run(para, display, bookmark_name)
            else:
                placeholder_run = para.add_run(f"[cite:{tag}]")
                set_font(placeholder_run, "宋体", 12)
                self._pending_cites.append({
                    'para_element': para._element,
                    'run_element': placeholder_run._element,
                    'tag': tag,
                })
            if idx < len(tags) - 1:
                run = para.add_run("、")
                set_font(run, "宋体", 12)

    def _resolve_pending_cites_for_tag(self, tag):
        """解析指定标签的所有待处理引用（替换占位符为超链接）"""
        if tag not in self._cite_registry:
            return
        ref_type, ref_num = self._cite_registry[tag]
        bookmark_name = f"cite_{tag}"
        if ref_type == "table":
            display = f" {ref_num}"
        elif ref_type == "equation":
            display = f" ({ref_num})"
        elif ref_type == "figure":
            display = f" {ref_num}"
        elif ref_type == "code":
            display = f" {ref_num}"
        elif ref_type == "theorem":
            display = f" {ref_num}"
        else:
            display = f" {ref_num}"

        remaining = []
        for entry in self._pending_cites:
            if entry['tag'] == tag:
                para_element = entry['para_element']
                old_run = entry['run_element']
                # 构建超链接元素替代占位 run
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('w:anchor'), bookmark_name)
                hyperlink.set(qn('w:history'), '1')
                new_run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '24')
                rPr.append(sz)
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '000000')
                rPr.append(color)
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'none')
                rPr.append(u)
                new_run.append(rPr)
                t = OxmlElement('w:t')
                t.text = display
                new_run.append(t)
                hyperlink.append(new_run)
                para_element.replace(old_run, hyperlink)
            else:
                remaining.append(entry)
        self._pending_cites = remaining

    def set_info(self, **kwargs):
        """设置论文信息"""
        self.info.update(kwargs)

    def add_cover(self):
        """添加封面"""
        _add_cover(self.doc, self.info, self.type)


    def add_abstract_cn(self, abstract_text, keywords, add_to_toc=True):
        """添加中文摘要"""
        para_title, formatted_title = self._create_heading_para('摘要')

        paragraphs = abstract_text.split('\n\n')
        for i, text in enumerate(paragraphs):
            text = text.strip()
            if not text:
                continue
            para = self._make_body_para()
            self._add_text_with_superscripts(para, text)

        if keywords:
            para = self.doc.add_paragraph()
            para.paragraph_format.line_spacing = Pt(20.5)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run('关键词：')
            set_font(run, "黑体", 14, True)
            keywords_text = "；".join(keywords) if isinstance(keywords, list) else keywords
            run2 = para.add_run(keywords_text)
            set_font(run2, "宋体", 12)

        return None

    def add_abstract_en(self, abstract_text, keywords, add_to_toc=True):
        """添加英文摘要"""
        para_title = self.doc.add_paragraph()
        add_page_break_before(para_title)
        apply_heading_style(para_title, after_pt=15)  # 官方模板Abstract间距更小
        run = para_title.add_run('Abstract')
        set_font(run, "Times New Roman", 18, False)
        set_outline_level(para_title, 0)

        paragraphs = abstract_text.split('\n\n')
        for i, text in enumerate(paragraphs):
            text = text.strip()
            if not text:
                continue
            para = self._make_body_para()
            run = para.add_run(text)
            set_font(run, "Times New Roman", 12)

        if keywords:
            para = self.doc.add_paragraph()
            para.paragraph_format.line_spacing = Pt(20.5)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run('Key words: ')
            set_font(run, "Times New Roman", 12, True)
            keywords_text = ", ".join(keywords) if isinstance(keywords, list) else keywords
            run2 = para.add_run(keywords_text)
            set_font(run2, "Times New Roman", 12)

        # 注意：英文摘要末尾不再添加分页
        return None

    def add_denotation(self, items=None, add_to_toc=True, title=None):
        """添加物理量名称及符号表

        items: 两列列表，每行 [符号, 含义]，如 [["D", "轴承直径，mm"], ["B", "轴承宽度，mm"]]
              若为 None，则创建空白2列表格供用户填写
        add_to_toc: 是否添加到目录
        title: 表题，如 "表1　国际单位制中具有专门名称的导出单位"（序号与名称之间空一格，即一个全角空格），
               若为 None 则使用默认表题
        """
        # 标题（与摘要同级格式：黑体18pt居中，自动分页）
        para = self.doc.add_paragraph()
        add_page_break_before(para)
        apply_heading_style(para)
        run = para.add_run('物理量名称及符号表')
        set_font(run, "黑体", 18, False)
        set_outline_level(para, 0)

        # 表题（宋体小四 14pt，居中）
        if title is None:
            title = "表1　主要物理量符号及说明"
        cap_para = self.doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.line_spacing = Pt(20.5)
        cap_para.paragraph_format.space_before = Pt(0)
        cap_para.paragraph_format.space_after = Pt(6)
        cap_run = cap_para.add_run(title)
        set_font(cap_run, "宋体", 10.5, False)

        # 三线表：2列（符号 | 含义），行数由 items 决定
        rows = 1 if items is None else max(len(items), 1)
        t = Table(self, rows, 2, caption=None, label=None)
        table_obj = t.insert()

        if items is not None:
            for i, (symbol, meaning) in enumerate(items):
                table_obj.set_cell(i, 0, symbol, bold=False)
                table_obj.set_cell(i, 1, meaning, bold=False)
        else:
            # 空白表格，提示用户填写
            table_obj.set_cell(0, 0, "", bold=False)
            table_obj.set_cell(0, 1, "", bold=False)

        return table_obj

    def add_toc(self, blank_line_before=False):
        """添加目录"""
        self._toc_blank_line = blank_line_before
        # 目录标题
        para = self.doc.add_paragraph()
        add_page_break_before(para)
        apply_heading_style(para)
        set_outline_level(para, 0)  # 设置大纲级别，使目录标题出现在 Word 导航栏
        run = para.add_run("目　录")
        set_font(run, "黑体", 18, False)

        # TOC 域代码：\o "1-3" 收集1-3级标题，\h 超链接，\z 隐藏Web标签，\u 使用大纲级别
        # 打开文档后 Word 需要更新此域才能显示目录内容（Windows 编译时自动更新）
        para = self.doc.add_paragraph()
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = r' TOC \o "1-3" \h \z \u '
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

    def add_page_break(self):
        """普通分页（不创建新节）"""
        self.doc.add_page_break()

    def add_authorization(self, add_to_toc=True):
        """添加授权声明"""
        _add_authorization(self.doc, self.info, add_to_toc, thesis_type=self.type)
    def add_conclusion(self, title='结论', content=None, intro=None, add_to_toc=True, auto_space=True):
        """添加结论
        结论作为论文正文的组成部分，单独排写，不加章标题序号。
        intro: 可选的导语段落（不编号），如 "本文取得了以下主要成果："
        content: 每条自动编号为（1）（2）...
        auto_space: 两汉字标题自动加全角空格
        """
        para, formatted_title = self._create_heading_para(title, auto_space=auto_space)
        if intro:
            para_intro = self._make_body_para()
            self._add_text_with_superscripts(para_intro, intro)
        if content:
            for i, text in enumerate(content):
                para2 = self._make_body_para()
                disable_snap_to_grid(para2)
                prefix = f"（{i + 1}）"
                self._add_text_with_superscripts(para2, f"{prefix}{text}")

    def add_acknowledgements(self, title='致谢', content=None, add_to_toc=True, auto_space=True):
        """添加致谢
        auto_space: 两汉字标题自动加全角空格
        """
        para, formatted_title = self._create_heading_para(title, auto_space=auto_space)
        if content:
            for text in content:
                para2 = self._make_body_para()
                self._add_text_with_superscripts(para2, text)

    def add_publications(self, title='攻读博士学位期间发表的论文及其他成果', sections=None, add_to_toc=True, auto_space=True):
        """添加发表文章页

        Args:
            title: 页面标题
            sections: [(副标题, [条目文字, ...]), ...]
            add_to_toc: 是否加入目录
            auto_space: 两汉字标题自动加全角空格
        """
        para, formatted_title = self._create_heading_para(title, auto_space=auto_space)
        if sections:
            for sub_title, items in sections:
                sub_para = self.doc.add_paragraph()
                sub_para.paragraph_format.line_spacing = Pt(20.5)
                sub_para.paragraph_format.space_before = Pt(12)
                sub_para.paragraph_format.space_after = Pt(6)
                run = sub_para.add_run(sub_title)
                set_font(run, "宋体", 12, True)

                for idx, item in enumerate(items):
                    item_para = self.doc.add_paragraph()
                    item_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    item_para.paragraph_format.line_spacing = Pt(20.5)
                    item_para.paragraph_format.space_before = Pt(0)
                    item_para.paragraph_format.space_after = Pt(0)
                    set_hanging_indent(item_para, 450)
                    run_idx = item_para.add_run(f"[{idx + 1}] ")
                    set_font(run_idx, "Times New Roman", 12)
                    self._add_text_with_superscripts(item_para, item)

    def add_resume(self, title='个人简历', content=None, add_to_toc=True, auto_space=True):
        """添加个人简历

        Args:
            title: 页面标题
            content: 正文段落字符串列表
            add_to_toc: 是否加入目录
            auto_space: 两汉字标题自动加全角空格
        """
        para, formatted_title = self._create_heading_para(title, auto_space=auto_space)
        if content:
            for text in content:
                para2 = self._make_body_para()
                self._add_text_with_superscripts(para2, text)

    def add_references(self, bib=None, title='参考文献', add_to_toc=True, references=None):
        """添加参考文献

        参数:
            bib: ReferenceDB 实例，或普通字符串列表（兼容旧用法）
            title: 参考文献标题
            add_to_toc: 是否加入目录
            references: 兼容旧用法的关键字参数，等同于 bib
        """
        para = self.doc.add_paragraph()
        add_page_break_before(para)
        apply_heading_style(para)
        run = para.add_run(title)
        set_font(run, '黑体', 18, False)
        set_outline_level(para, 0)

        # 兼容旧用法：优先使用 references 参数
        if bib is None and references is not None:
            bib = references

        if bib is None:
            return

        # 兼容字符串列表（旧用法）
        if isinstance(bib, list):
            for ref in bib:
                para2 = self.doc.add_paragraph()
                para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para2.paragraph_format.line_spacing = Pt(20.5)
                para2.paragraph_format.space_before = Pt(0)
                para2.paragraph_format.space_after = Pt(6)
                set_hanging_indent(para2, 450)  # 悬挂缩进对齐[1] 后文字
                run2 = para2.add_run(ref)
                set_font(run2, 'Times New Roman', 12)
            return

        # ReferenceDB 用法：自动按引用顺序生成
        from hitthesis.reference_db import ReferenceDB
        if not isinstance(bib, ReferenceDB):
            return

        # 按引用顺序获取格式化文献，添加编号和书签
        refs = bib.get_references_in_citation_order()
        for idx, ref_str in enumerate(refs, start=1):
            para2 = self.doc.add_paragraph()
            para2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para2.paragraph_format.line_spacing = Pt(20.5)
            para2.paragraph_format.space_before = Pt(0)
            para2.paragraph_format.space_after = Pt(6)
            disable_snap_to_grid(para2)
            set_hanging_indent(para2, 450)

            # 编号部分 + 文献内容合成一个 run
            run = para2.add_run(f"[{idx}] {ref_str}")
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            rPr.append(rFonts)
            run.font.size = Pt(12)

            # 为段落添加书签（供超链接跳转）
            bookmark_name = f"ref_{idx}"
            add_bookmark_to_paragraph(para2, bookmark_name)

    def add_appendix(self, title='附录', add_to_toc=True):
        """添加附录（返回上下文管理器，支持add_section等方法）
        用法：
            with doc.add_appendix("符号表"):
                doc.add_section("主要符号")
                doc.add_subsection("几何参数")
        """
        self.appendix_number += 1
        # 设置chapter_number为附录号，section_counter从0开始
        self.chapter_number = self.appendix_number
        self.section_counter = [0, 0, 0]
        ctx = AppendixContext(self, title, self.appendix_number, add_to_toc)
        return ctx

    def include(self, module_path):
        """导入外部章节模块并执行其 build(doc) 函数。

        用法:
            doc.include("body.chapter1_intro")

        module_path 为 Python 模块路径（点号分隔），对应文件系统 body/chapter1_intro.py。
        章节文件需导出 build(doc) 函数，接收 Thesis 实例作为唯一参数。
        需确保模块所在目录在 sys.path 中。
        """
        import importlib
        mod = importlib.import_module(module_path)
        mod.build(self)

    def compile(self, filename):
        """编译生成论文文档"""
        if self._pending_cites:
            unresolved = set(entry['tag'] for entry in self._pending_cites)
            print(f"警告: 以下引用标签未注册，将以原始文本保留: {unresolved}")
        compile_document(self.doc, filename, thesis_type=self.type,
                         toc_blank_line=self._toc_blank_line,
                         footnotes=self._footnotes if self._footnotes else None)

    def _format_title(self, title, with_chapter_prefix=None, auto_space=True):
        """格式化一级标题：
        - 检测两汉字则在中间插入全角空格（结　论），auto_space 可关闭
        - 可选添加"第X章　"前缀
        """
        import re
        if auto_space and len(title) == 2 and re.match(r'^[一-龥]+$', title):
            title = f"{title[0]}　{title[1]}"
        if with_chapter_prefix is not None:
            title = f"第{with_chapter_prefix}章　{title}"
        return title

    def _create_heading_para(self, title, with_chapter_prefix=None, auto_space=True, after_pt=None):
        """创建一级标题段落（复用逻辑）
        with_chapter_prefix: 指定章号，如"1"，自动生成"第1章　标题"
        auto_space: 两汉字标题自动加全角空格
        after_pt: 自定义段后间距（pt），None则使用默认值
        返回: (段落, 格式化后的标题文字)
        """
        para = self.doc.add_paragraph()
        add_page_break_before(para)
        if after_pt is not None:
            apply_heading_style(para, after_pt=after_pt)
        else:
            apply_heading_style(para)
        formatted_title = self._format_title(title, with_chapter_prefix, auto_space)
        run = para.add_run(formatted_title)
        set_font(run, '黑体', 18, False)
        set_outline_level(para, 0)
        return para, formatted_title

    def add_chapter(self, title, number=None, spacing_before=Pt(0)):
        """添加章节（返回上下文管理器，由with语句调用__enter__）
        spacing_before: 段前间距，默认0，第一章可设为Pt(24)实现目录后空一行
        """
        ctx = ChapterContext(self, title, number, spacing_before)  # 传递self（Thesis对象）
        return ctx

    def _generate_section_number(self, level):
        """生成小节编号
        level: 1=一级(section), 2=二级(subsection), 3=三级(subsubsection)
        """
        ch = self.chapter_number
        if level == 1:
            self.section_counter[0] += 1
            self.section_counter[1] = 0
            self.section_counter[2] = 0
            return f"{ch}.{self.section_counter[0]}"
        elif level == 2:
            self.section_counter[1] += 1
            self.section_counter[2] = 0
            return f"{ch}.{self.section_counter[0]}.{self.section_counter[1]}"
        elif level == 3:
            self.section_counter[2] += 1
            return f"{ch}.{self.section_counter[0]}.{self.section_counter[1]}.{self.section_counter[2]}"
        return ""

    def _format_section_title(self, title, level):
        """格式化小节标题：在编号和标题之间插入一个全角空格"""
        import re
        section_num = self._generate_section_number(level)
        return f"{section_num}　{title}"

    def add_section(self, title):
        """一级小节标题（section）"""
        formatted_title = self._format_section_title(title, 1)
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing = Pt(22)    # 21bp行距
        para.paragraph_format.space_before = Pt(SPACING["section_before"])
        para.paragraph_format.space_after = Pt(SPACING["section_after"])
        run = para.add_run(formatted_title)
        set_font(run, "黑体", 15, False)               # bachelor不加粗
        set_outline_level(para, 1)
        add_heading_properties(para)
        return para

    def add_subsection(self, title):
        """二级小节标题（subsection）"""
        formatted_title = self._format_section_title(title, 2)
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing = Pt(19)    # 18bp行距
        para.paragraph_format.space_before = Pt(SPACING["subsection_before"])
        para.paragraph_format.space_after = Pt(SPACING["subsection_after"])
        run = para.add_run(formatted_title)
        set_font(run, "黑体", 14, False)               # bachelor不加粗
        set_outline_level(para, 2)
        add_heading_properties(para)
        return para

    def add_subsubsection(self, title):
        """三级小节标题（subsubsection）——不加入目录"""
        formatted_title = self._format_section_title(title, 3)
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.line_spacing = Pt(16)    # 约12bp×1.33行距
        para.paragraph_format.space_before = Pt(SPACING["subsubsection_before"])
        para.paragraph_format.space_after = Pt(SPACING["subsubsection_after"])
        run = para.add_run(formatted_title)
        set_font(run, "黑体", 12, False)               # normalsize 12bp，不加粗
        add_heading_properties(para)
        return para





    def _make_body_para(self, indent=True):
        """创建标准正文段落：两端对齐、1.3倍行距、段前后为0、可选首行缩进"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = Pt(SPACING["body_line_spacing"])
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if indent:
            set_first_line_indent(para, SPACING["first_line_indent"])
        return para

    def add_paragraph(self, text, indent=True, footnote=None, footnote_number=None):
        """添加正文段落（自动识别[数字]引用并设为上标）

        Args:
            text: 段落文字
            indent: 是否首行缩进
            footnote: 可选，脚注文字。自动在段落末尾插入上标标记并在本页底部显示脚注
            footnote_number: 可选，自定义脚注序号。None 时使用全局递增序号
        """
        para = self._make_body_para(indent=indent)
        self._add_text_with_superscripts(para, text)
        if footnote:
            self._insert_footnote_marker(para, footnote, number=footnote_number)
        return para

    def _insert_footnote_marker(self, para, text, number=None):
        """在段落末尾标点符号前插入脚注上标标记。

        脚注标记需要插在标点前面（如"研究。①"而非"研究①。"），所以要拆分末尾标点：
        正文文字留在原 run → 标记 run → 标点移到新 run。

        Args:
            para: 段落对象
            text: 脚注文字
            number: 可选，自定义脚注序号。None 时使用全局递增序号
        """
        if number is None:
            self._footnote_counter += 1
            display_num = self._footnote_counter
        else:
            self._footnote_counter += 1
            display_num = number
        global_id = self._footnote_counter
        self._footnotes.append((global_id, display_num, text))
        CIRCLE = "①②③④⑤⑥⑦⑧⑨"
        marker = CIRCLE[display_num - 1] if display_num <= 9 else f"[{display_num}]"

        # 创建隐藏 run：w:footnoteReference（Word 内部脚注定位用，1pt 不可见）
        def _make_footnote_ref_run():
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'FootnoteReference'); rPr.append(rStyle)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '2'); rPr.append(sz)
            r.append(rPr)
            fn_ref = OxmlElement('w:footnoteReference')
            fn_ref.set(qn('w:id'), str(global_id))
            r.append(fn_ref)
            return r

        # 创建可见标记：包裹在 w:hyperlink 中，点击跳转到 footnotes.xml 的书签 _ftn{idx}
        def _make_marker_hyperlink():
            hl = OxmlElement('w:hyperlink')
            hl.set(qn('w:anchor'), f'_ftn{global_id}')
            hl.set(qn('w:history'), '1')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rPr.append(sz)
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:eastAsia'), '宋体'); rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rf)
            va = OxmlElement('w:vertAlign'); va.set(qn('w:val'), 'superscript'); rPr.append(va)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = marker; r.append(t)
            hl.append(r)
            return hl

        W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        runs = para._element.findall(f'{{{W}}}r')
        punct_trail = "。！？）』」》〉\"'.;!?)"

        if not runs:
            para._element.append(_make_footnote_ref_run())
            para._element.append(_make_marker_hyperlink())
            return

        last_r = runs[-1]
        last_t = last_r.find(f'{{{W}}}t')
        if last_t is None or not last_t.text:
            para._element.append(_make_footnote_ref_run())
            para._element.append(_make_marker_hyperlink())
            return

        text_end = last_t.text
        # 收集末尾连续标点
        trail = ""
        while text_end and text_end[-1] in punct_trail:
            trail = text_end[-1] + trail
            text_end = text_end[:-1]

        if trail:
            # 拆分：正文部分留在原 run，标点移到新 run，标记插中间
            last_t.text = text_end if text_end else None
            punct_r = OxmlElement('w:r')
            # 复制 rPr
            old_rPr = last_r.find(f'{{{W}}}rPr')
            if old_rPr is not None:
                punct_r.append(copy.deepcopy(old_rPr))
            pt = OxmlElement('w:t')
            pt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            pt.text = trail; punct_r.append(pt)
            para._element.append(_make_footnote_ref_run())
            para._element.append(_make_marker_hyperlink())
            para._element.append(punct_r)
        else:
            para._element.append(_make_footnote_ref_run())
            para._element.append(_make_marker_hyperlink())


    def set_page_number_format(self, fmt, start=None):
        """为当前节（最后一个节）设置页码格式
        fmt: 'I' 罗马数字, '1' 阿拉伯数字
        start: 起始页码（阿拉伯数字需要，罗马数字不需要）
        """
        section = self.doc.sections[-1]
        sectPr = section._sectPr
        # 移除已有的 pgNumType
        old = sectPr.find(qn('w:pgNumType'))
        if old is not None:
            sectPr.remove(old)
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:val'), fmt)
        if start is not None:
            pgNumType.set(qn('w:start'), str(start))
        sectPr.append(pgNumType)
        # 设置页眉页脚
        self._setup_header_footer_for_section(section)

    def start_roman_section(self):
        """开始罗马数字页码节（封面之后，摘要之前）"""
        section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        sectPr = section._sectPr
        old = sectPr.find(qn('w:pgNumType'))
        if old is not None:
            sectPr.remove(old)

        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), 'upperRoman')
        pgNumType.set(qn('w:start'), '1')
        sectPr.append(pgNumType)

        self._setup_header_footer_for_section(section)

    def start_arabic_section(self):
        """开始阿拉伯数字页码节（目录之后，正文第一章之前）"""
        section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        sectPr = section._sectPr
        old = sectPr.find(qn('w:pgNumType'))
        if old is not None:
            sectPr.remove(old)

        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), '1')
        sectPr.append(pgNumType)

        self._setup_header_footer_for_section(section)

    def add_table(self, rows, cols, caption=None, label=None, ref=None):
        """添加三线表（caption在上方，表格自动应用三线表样式）

        Args:
            rows: 行数
            cols: 列数
            caption: 表题
            label: 编号，如 "1-1"
            ref: 引用标签，如 "results"
        """
        # 自动编号
        if label:
            ref_num = label
        else:
            self._tab_counter += 1
            ref_num = f"{self.chapter_number}-{self._tab_counter}"

        # 注册引用
        if ref:
            self._cite_registry[ref] = ("table", ref_num)
            self._resolve_pending_cites_for_tag(ref)

        t = Table(self, rows, cols, caption=caption, label=ref_num, ref=ref)
        return t.insert()


    def add_figure(self, image_path=None, caption=None, width=None, label=None, ref=None):
        """添加图片

        Args:
            image_path: 图片路径
            caption: 图题
            width: 图片宽度
            label: 编号，如 "1-1"
            ref: 引用标签，如 "sample"
        """
        # 自动编号
        if label:
            ref_num = label
        else:
            self._fig_counter += 1
            ref_num = f"{self.chapter_number}-{self._fig_counter}"

        # 注册引用
        if ref:
            self._cite_registry[ref] = ("figure", ref_num)
            self._resolve_pending_cites_for_tag(ref)

        fig = Figure(self, image_path=image_path, caption=caption, width=width, label=ref_num, ref=ref)
        return fig.insert()

    def add_subfigure(self, subfigures, caption=None, label=None, ref=None, cols=None):
        """添加子图（支持并排或网格布局）

        Args:
            subfigures: [(image_path, subcaption), ...] 子图列表
            caption: 总图题
            label: 编号，如 "1-1"
            ref: 引用标签
            cols: 每行列数，默认自动（1行并排）
        """
        # 自动编号
        if label:
            ref_num = label
        else:
            self._fig_counter += 1
            ref_num = f"{self.chapter_number}-{self._fig_counter}"

        # 注册引用
        if ref:
            self._cite_registry[ref] = ("figure", ref_num)
            self._resolve_pending_cites_for_tag(ref)

        fig = SubFigure(self, subfigures, caption=caption, label=ref_num, ref=ref, cols=cols)
        return fig.insert()

    def add_code_block(self, code_text, caption=None, label=None, ref=None):
        """添加代码块（浅灰底纹、等宽字体、题注在上方）

        Args:
            code_text: 代码文本，换行用 \\n
            caption: 代码题注
            label: 编号，如 "2-1"
            ref: 引用标签，如 "quick_sort"
        """
        if label:
            ref_num = label
        else:
            self._code_counter += 1
            ref_num = f"{self.chapter_number}-{self._code_counter}"

        if ref:
            self._cite_registry[ref] = ("code", ref_num)
            self._resolve_pending_cites_for_tag(ref)

        # ── 题注在上方 ──
        if caption:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.line_spacing = Pt(14)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            # 附录内 code caption 加"附"前缀：附代码 1-1
            code_word = "附代码" if self._in_appendix else "代码"
            text = f"{code_word}{ref_num}　{caption}"

            add_caption_with_bookmark(para, ref, text)

        # ── 代码体 ──
        CODE_GRAY = "F5F5F5"
        for line in code_text.split('\n'):
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = Pt(14)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            # 整体左缩进 0.5cm (283 twips ≈ 0.5cm)
            para.paragraph_format.left_indent = Cm(0.5)
            # 无首行缩进
            para.paragraph_format.first_line_indent = Pt(0)
            disable_snap_to_grid(para)

            # 浅灰底纹
            pPr = para._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), CODE_GRAY)
            pPr.append(shd)

            run = para.add_run(line if line else ' ')
            # Consolas 是西文字体，需同时设 ascii/hAnsi/eastAsia
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.bold = False
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), 'Consolas')
            rFonts.set(qn('w:hAnsi'), 'Consolas')
            rFonts.set(qn('w:eastAsia'), 'Consolas')

    def add_equation(self, formula, label=None, ref=None):
        """插入公式（无边框表格，居中，1.5倍行距）

        Args:
            formula: LaTeX 公式文本
            label: 编号字符串，格式 "章号-序号"，如 "2-1"
            ref: 可选，书签名称，如 "eq_2_1"
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        # 生成编号
        if label:
            number_text = f"({label})"
            ref_num = label  # 去掉括号
        else:
            self._eq_counter += 1
            ref_num = f"{self.chapter_number}-{self._eq_counter}"
            # 附录内 number_text 加"附"前缀：(附1-1)；章节内：(1-1)
            display_num = f"附{ref_num}" if self._in_appendix else ref_num
            number_text = f"({display_num})"

        # 注册引用
        if ref:
            self._cite_registry[ref] = ("equation", ref_num)
            self._resolve_pending_cites_for_tag(ref)

        # 获取正文宽度
        section = self.doc.sections[-1]
        page_width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
        col1_width = int(page_width * 0.15)
        col3_width = int(page_width * 0.15)
        col2_width = page_width - col1_width - col3_width

        # 创建表格（1行×3列）
        table = self.doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表格宽度100%
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(page_width))  # 正文总宽度
        tblW.set(qn('w:type'), 'dxa')  # 固定值
        tblPr.append(tblW)

        # 去除所有边框
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        cells = table.rows[0].cells

        # 第1列：留白（15%）
        self._setup_equation_cell(cells[0], "", 'left', 'center', col1_width)

        # 第2列：公式文本（居中，1.5倍行距，70%）
        para2_elem = self._setup_equation_cell(cells[1], formula, 'center', 'center', col2_width)
        pPr2 = para2_elem.get_or_add_pPr()
        spacing2 = OxmlElement('w:spacing')
        spacing2.set(qn('w:line'), '360')  # 1.5倍行距
        spacing2.set(qn('w:lineRule'), 'auto')
        spacing2.set(qn('w:before'), '152')  # 段前
        spacing2.set(qn('w:after'), '152')   # 段后
        pPr2.append(spacing2)

        # 第3列：编号（右对齐，垂直居中，15%）
        para3_elem = self._setup_equation_cell(cells[2], number_text, 'right', 'center', col3_width)
        pPr3 = para3_elem.get_or_add_pPr()
        spacing3 = OxmlElement('w:spacing')
        spacing3.set(qn('w:line'), '360')
        spacing3.set(qn('w:lineRule'), 'auto')
        spacing3.set(qn('w:before'), '152')
        spacing3.set(qn('w:after'), '152')
        pPr3.append(spacing3)
        # 添加书签（交叉引用用）—— 必须放在段落级 <w:p> 内，不能放 <w:tc>
        if ref:
            add_bookmark(para3_elem, f"cite_{ref}")

        return self

    def _setup_equation_cell(self, cell, text, h_align, v_align, width_twips=None):
        """配置公式单元格（无边框，指定对齐）"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # 设置单元格宽度
        if width_twips is not None:
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width_twips))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        # 无边框
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tcBorders.append(border)
        tcPr.append(tcBorders)

        # 垂直居中
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), v_align)
        tcPr.append(vAlign)

        # 段落
        para = cell.paragraphs[0]
        if h_align == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif h_align == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 公式渲染：如果是 LaTeX 公式，调用 latex2word
        if text and '\\' in text:
            from .latex_render import add_latex_to_paragraph
            add_latex_to_paragraph(para, text, set_font)
        else:
            run = para.add_run(text or "")
            set_font(run, "Times New Roman", 12)

        # 返回段落元素（用于在段落级插入 bookmark）
        return para._element


    def add_footnote(self, text, number=None):
        """插入脚注标记到上一个段落末尾（向后兼容）

        Args:
            text: 脚注文字
            number: 可选，自定义脚注序号。None 时使用全局递增序号
        """
        paras = self.doc.paragraphs
        target_para = paras[-1] if paras else self.doc.add_paragraph()
        self._insert_footnote_marker(target_para, text, number=number)

    def add_page_break(self):
        """手动分页"""
        para = self.doc.add_paragraph()
        run = para.add_run()
        run._element.append(OxmlElement('w:br'))
        run._element[-1].set(qn('w:type'), 'page')

    def add_quote(self, text):
        """引用块：楷体 11pt，左右缩进，段前后 3pt"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = Pt(18)
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.left_indent = Cm(1.0)
        para.paragraph_format.right_indent = Cm(1.0)
        para.paragraph_format.first_line_indent = Pt(0)
        disable_snap_to_grid(para)
        run = para.add_run(text)
        run.font.name = '楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        run.font.size = Pt(SPACING["body_font_size"])  # 12pt，与正文一致
        return para

    def add_list(self, items, style="decimal"):
        """有序/无序列表：前缀顶格，无悬挂缩进

        Args:
            items: 列表项文本列表
            style: "decimal"=1. 2. 3.  "paren"=（1）（2）  "bullet"=—
        """
        for i, item in enumerate(items):
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing = Pt(SPACING["body_line_spacing"])
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1)
            set_first_line_indent(para, SPACING["first_line_indent"])
            disable_snap_to_grid(para)
            if style == "paren":
                prefix = f"（{i + 1}）"
            elif style == "bullet":
                prefix = "—"
            else:
                prefix = f"{i + 1}."
            run = para.add_run(f"{prefix}  {item}" if style != "paren" else f"{prefix}{item}")
            set_font(run, "宋体", SPACING["body_font_size"], False)

    def add_theorem(self, text, kind="定理", label=None, ref=None, cite=None):
        """定理/定义/引理环境：黑体标题 + 宋体正文同行，自动编号（章.序号）

        Args:
            text: 定理内容
            kind: 类型 "定理" "定义" "引理" "推论" 等
            label: 手动编号
            ref: 交叉引用标签
            cite: 可选标注，如 "[ref:Einstein]" 或 "Albert Einstein"
        """
        if label:
            ref_num = label
        else:
            self._thm_counter += 1
            ref_num = f"{self.chapter_number}.{self._thm_counter}"
        if ref:
            self._cite_registry[ref] = ("theorem", ref_num)
            self._resolve_pending_cites_for_tag(ref)

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = Pt(SPACING["body_line_spacing"])
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        set_first_line_indent(para, SPACING["first_line_indent"])
        disable_snap_to_grid(para)

        run = para.add_run(f"{kind} {ref_num}")
        set_font(run, "黑体", SPACING["body_font_size"], True)
        if ref:
            add_bookmark(para._element, f"cite_{ref}")
        if cite:
            self._add_text_with_superscripts(para, f"{cite}")
        self._add_text_with_superscripts(para, f"  {text}" if text else "")




