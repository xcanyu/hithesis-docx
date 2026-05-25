r"""
LaTeX 公式与化学式渲染模块
触发方式：文本中用 $...$ 包裹
规则：
  - ^x  单字符上标
  - ^{xyz} 多字符上标
  - _x  单字符下标
  - _{xyz} 多字符下标
  - 其他 LaTeX 命令（如 \frac, \sqrt）调用 latex2word 处理
"""

import re
import copy
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _parse_latex_with_scripts(latex_str: str) -> list:
    """解析 LaTeX 内容，处理上下标。

    返回: [(text, style), ...]
    style: 'normal' | 'superscript' | 'subscript'

    例如 "SO4^2-" -> [('S','normal'),('O','normal'),('4','subscript'),('2-','superscript')]
    """
    result = []
    i = 0
    n = len(latex_str)

    while i < n:
        char = latex_str[i]

        if char == '^':
            # 上标：收集 ^ 后的内容，直到非字母数字字符为止
            if i + 1 < n:
                next_char = latex_str[i + 1]
                if next_char == '{':
                    # 找匹配的 }
                    end = latex_str.find('}', i + 1)
                    if end != -1:
                        content = latex_str[i + 2:end]
                        result.append((content, 'superscript'))
                        i = end + 1
                        continue
                else:
                    # 多字符上标：收集连续的非特殊分隔字符
                    j = i + 2
                    while j < n and latex_str[j] not in (' ', '\t', '+', '-'):
                        j += 1
                    # 末尾的 - 也纳入上标
                    if j < n and latex_str[j] == '-':
                        j += 1
                    content = latex_str[i + 1:j]
                    result.append((content, 'superscript'))
                    i = j
                    continue
            i += 1
            continue

        if char == '_':
            # 下标
            if i + 1 < n:
                next_char = latex_str[i + 1]
                if next_char == '{':
                    end = latex_str.find('}', i + 1)
                    if end != -1:
                        content = latex_str[i + 2:end]
                        result.append((content, 'subscript'))
                        i = end + 1
                        continue
                else:
                    # 多字符下标
                    j = i + 2
                    while j < n and latex_str[j].isalnum():
                        j += 1
                    content = latex_str[i + 1:j]
                    result.append((content, 'subscript'))
                    i = j
                    continue
            i += 1
            continue

        # 普通字符
        result.append((char, 'normal'))
        i += 1

    return result


def needs_latex2word(latex_str: str) -> bool:
    """判断是否需要调用 latex2word 处理（包含复杂 LaTeX 命令）"""
    # 简单检查：是否包含 \ 开头的 LaTeX 命令
    return '\\' in latex_str


def add_latex_to_paragraph(para, latex_str: str, font_func):
    """向段落添加 LaTeX 公式或化学式

    Args:
        para: docx Paragraph 对象
        latex_str: LaTeX 字符串，不含外层 $
        font_func: 设置字体的函数，接收 (run, font_name, size, bold)
    """
    if needs_latex2word(latex_str):
        # 复杂 LaTeX，调用 latex2word
        from docx import Document
        temp_doc = Document()
        temp_doc.add_paragraph()  # 先添加一个段落供 latex2word 使用
        try:
            from latex2word import LatexToWordElement
            converter = LatexToWordElement(latex_str)
            converter.add_latex_to_paragraph(temp_doc.paragraphs[0])
            # 从临时段落复制内容到目标段落（使用 deep copy 避免命名空间问题）
            src_p = temp_doc.paragraphs[0]._p
            for child in src_p:
                para._p.append(copy.deepcopy(child))
        except Exception as e:
            # 转换失败，回退到显示原文本
            run = para.add_run(f"${latex_str}$")
            font_func(run, "Times New Roman", 12, False)
    else:
        # 简单上下标处理
        parts = _parse_latex_with_scripts(latex_str)
        for text, style in parts:
            run = para.add_run(text)
            font_func(run, "Times New Roman", 12, False)
            if style == 'superscript':
                run.font.superscript = True
            elif style == 'subscript':
                run.font.subscript = True


def split_text_with_latex(text: str) -> list:
    """将文本分割为普通文本和 LaTeX 片段

    返回: [('text', '普通文本'), ('latex', 'LaTeX内容'), ...]
    """
    parts = []
    pattern = re.compile(r'\$([^$]+)\$')
    last_end = 0

    for m in pattern.finditer(text):
        if m.start() > last_end:
            parts.append(('text', text[last_end:m.start()]))
        parts.append(('latex', m.group(1)))
        last_end = m.end()

    if last_end < len(text):
        parts.append(('text', text[last_end:]))

    return parts